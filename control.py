from models import Exam, ExamResult, Question, db, Link, KnowledgePoint, DownloadRecord, User, Class, ExamClassAssociation, Chapter, Section, Material, Post, Reply, Blog, SignInRecord, RouteLog, LearningMaterial,LearningMaterialClass, ExamSession, Notice, RecommendedReading, chongbuluo_URL,HTMLPage, AISession, AIMessage, AIPromptTemplate, AIModelConfig, AIApiKey
from config import app, photos, handler, logger, load_config
from utils import calculate_similarity, secure_filename_with_chinese, convert_relative_paths_to_absolute, chatgpt, call_baidu_gpt_api, call_xunfei_gpt
from werkzeug.utils import secure_filename
from flask import Flask, request, jsonify, render_template, send_from_directory, redirect, url_for, flash, Response, send_file, abort, stream_with_context, after_this_request
from flask_login import login_required, current_user, LoginManager, login_user, UserMixin, logout_user
from flask_uploads import UploadSet, configure_uploads
from sqlalchemy import or_, desc
from sqlalchemy.orm import aliased
from sqlalchemy.sql import exists
from ai_service import AIService
from users import permissions_required
import os
from datetime import date, datetime, timedelta
import json
import pandas as pd
from openpyxl import load_workbook
from flask_sqlalchemy import SQLAlchemy
from bs4 import BeautifulSoup, Comment
from docx import Document
import random
import pypandoc
import io
import time  # 引入time用于调试

import logging

# 配置日志记录
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 错误处理
@app.errorhandler(403)
def forbidden(e):
    return render_template('403.html'), 403

#测试，跨网站登录
@app.route('/hydrooj-login')
def simple_login():
    """简化的自动登录版本"""
    
    # 目标服务器信息
    TARGET_SERVER = "http://2.75.209.63/d/xinxixuexuankao2028"
    
    username = current_user.username
    # 直接创建一个自动提交表单的页面
    html = f'''
    <!DOCTYPE html>
    <html>
    <head>
        <title>自动跳转到8081端口</title>
        <script>
        function autoLogin() {{
            // 创建隐藏表单
            var form = document.createElement('form');
            form.method = 'POST';
            form.action = '{TARGET_SERVER}/login';
            // 添加登录信息
            var credentials = {{
                'uname': '{username}',
                'password': '123456',
                'tfa': '',
                'authnChallenge': ''
            }};
            
            for (var key in credentials) {{
                var input = document.createElement('input');
                input.type = 'hidden';
                input.name = key;
                input.value = credentials[key];
                form.appendChild(input);
            }}
            
            document.body.appendChild(form);
            form.submit();
        }}
        
        // 页面加载后立即执行
        window.onload = autoLogin;
        </script>
    </head>
    <body>
        <p>正在自动跳转到8081端口并登录...</p>
        <p>如果页面没有自动跳转，<a href="#" onclick="autoLogin(); return false;">点击这里</a></p>
    </body>
    </html>
    '''
    return html

#AI prompt项目-----------------------------------------------------------------------------------
ai_service = AIService()

@app.route('/admin/ai-config')
@login_required
@permissions_required('is_admin')
def ai_config_management():
    """AI配置管理页面"""
    # 检查管理员权限（根据您的权限系统调整）
    if not current_user.is_admin:  # 假设您的User模型有is_admin字段
        return "无权限访问", 403
    return render_template('ai_config_management.html')

# API路由 - Prompt模板配置
@app.route('/api/admin/ai/templates', methods=['GET', 'POST', 'PUT', 'DELETE'])
@login_required
@permissions_required('is_admin')
def admin_ai_templates():
    """Prompt模板配置管理API"""
    if not current_user.is_admin:
        return jsonify({'error': '无权限'}), 403
    
    if request.method == 'GET':
        templates = AIPromptTemplate.query.all()
        templates_data = []
        for template in templates:
            templates_data.append({
                'id': template.id,
                'name': template.name,
                'content': template.content,
                'description': template.description,
                'variables': template.variables or [],
                'is_active': template.is_active,
                'created_at': template.created_at.isoformat()
            })
        return jsonify(templates_data)
    
    elif request.method == 'POST':
        data = request.get_json()
        new_template = AIPromptTemplate(
            name=data['name'],
            content=data['content'],
            description=data.get('description', ''),
            variables=data.get('variables', []),
            is_active=data.get('is_active', True)
        )
        db.session.add(new_template)
        db.session.commit()
        return jsonify({'success': True, 'id': new_template.id})
    
    elif request.method == 'PUT':
        data = request.get_json()
        template_id = data.get('id')
        template = AIPromptTemplate.query.get_or_404(template_id)
        
        template.name = data['name']
        template.content = data['content']
        template.description = data.get('description', '')
        template.variables = data.get('variables', [])
        template.is_active = data.get('is_active', True)
        
        db.session.commit()
        return jsonify({'success': True})
    
    elif request.method == 'DELETE':
        template_id = request.args.get('id')
        template = AIPromptTemplate.query.get_or_404(template_id)
        db.session.delete(template)
        db.session.commit()
        return jsonify({'success': True})

# API路由 - API密钥配置
@app.route('/api/admin/ai/api-keys', methods=['GET', 'POST', 'PUT', 'DELETE'])
@login_required
@permissions_required('is_admin')
def admin_ai_api_keys():
    """API密钥配置管理API"""
    if not current_user.is_admin:
        return jsonify({'error': '无权限'}), 403
    
    if request.method == 'GET':
        keys = AIApiKey.query.all()
        keys_data = []
        for key in keys:
            # 不返回完整的API密钥，只显示部分
            masked_key = key.api_key[:8] + '***' + key.api_key[-4:] if len(key.api_key) > 12 else '***'
            keys_data.append({
                'id': key.id,
                'provider': key.provider,
                'api_key_masked': masked_key,
                'is_active': key.is_active,
                'created_at': key.created_at.isoformat()
            })
        return jsonify(keys_data)
    
    elif request.method == 'POST':
        data = request.get_json()
        new_key = AIApiKey(
            provider=data['provider'],
            api_key=data['api_key'],
            is_active=data.get('is_active', True)
        )
        db.session.add(new_key)
        db.session.commit()
        return jsonify({'success': True, 'id': new_key.id})
    
    elif request.method == 'PUT':
        data = request.get_json()
        key_id = data.get('id')
        key_obj = AIApiKey.query.get_or_404(key_id)
        
        key_obj.provider = data['provider']
        if data.get('api_key'):  # 只有当提供了新密钥时才更新
            key_obj.api_key = data['api_key']
        key_obj.is_active = data.get('is_active', True)
        
        db.session.commit()
        return jsonify({'success': True})
    
    elif request.method == 'DELETE':
        key_id = request.args.get('id')
        key_obj = AIApiKey.query.get_or_404(key_id)
        db.session.delete(key_obj)
        db.session.commit()
        return jsonify({'success': True})

@app.route('/ai-chat')
@login_required
def ai_chat():
    """AI聊天主界面"""
    return render_template('ai_chat.html')

# API路由
@app.route('/api/ai/sessions', methods=['GET', 'POST'])
@login_required
def ai_sessions():
    """会话管理"""
    if request.method == 'GET':
        # 获取用户的所有会话
        sessions = AISession.query.filter_by(user_id=current_user.id)\
            .order_by(AISession.updated_at.desc()).all()
        
        sessions_data = []
        for s in sessions:
            sessions_data.append({
                'id': s.id,
                'title': s.title,
                'model_used': s.model_used,
                'prompt_template': s.prompt_template,
                'created_at': s.created_at.isoformat(),
                'updated_at': s.updated_at.isoformat(),
                'message_count': s.messages.count()
            })
        
        return jsonify(sessions_data)
    
    else:  # POST
        data = request.get_json()
        new_session = AISession(
            user_id=current_user.id,
            title=data.get('title', '新对话'),
            model_used=data.get('model_used', 'GPT-3.5-Turbo')
        )
        db.session.add(new_session)
        db.session.commit()
        
        return jsonify({
            'id': new_session.id,
            'title': new_session.title,
            'model_used': new_session.model_used
        })

@app.route('/api/ai/sessions/<int:session_id>', methods=['GET', 'PUT', 'DELETE'])
@login_required
def ai_session_detail(session_id):
    """会话详情、更新和删除"""
    ai_session = AISession.query.filter_by(id=session_id, user_id=current_user.id).first_or_404()

    if request.method == 'DELETE':
        db.session.delete(ai_session)
        db.session.commit()
        return jsonify({'success': True})

    elif request.method == 'PUT':
        data = request.get_json() or {}
        model_used = data.get('model_used')
        title = data.get('title')
        if model_used:
            ai_session.model_used = model_used
        if title:
            ai_session.title = title
        db.session.commit()
        return jsonify({'success': True})

    else:  # GET
        messages = ai_session.messages.order_by(AIMessage.created_at.asc()).all()
        messages_data = [{
            'role': msg.role,
            'content': msg.content,
            'created_at': msg.created_at.isoformat()
        } for msg in messages]

        return jsonify({
            'session': {
                'id': ai_session.id,
                'title': ai_session.title,
                'model_used': ai_session.model_used
            },
            'messages': messages_data
        })

@app.route('/api/ai/sessions/<int:session_id>/messages', methods=['POST'])
@login_required
def send_ai_message(session_id):
    """发送消息"""
    ai_session = AISession.query.filter_by(id=session_id, user_id=current_user.id).first_or_404()
    data = request.get_json()
    
    user_message_content = data.get('message', '').strip()
    template_id = data.get('template_id')
    template_variables = data.get('variables', {})
    
    if not user_message_content:
        return jsonify({'error': '消息内容不能为空'}), 400
    
    # 保存用户消息
    user_message = AIMessage(
        session_id=session_id,
        role='user',
        content=user_message_content,
        tokens=ai_service.estimate_tokens(user_message_content)
    )
    db.session.add(user_message)
    
    # 更新会话标题（如果是第一条消息）
    if ai_session.messages.count() == 1:  # 刚创建的用户消息是第一条
        ai_session.title = user_message_content[:50] + '...' if len(user_message_content) > 50 else user_message_content
    
    db.session.commit()
    
    # 获取模板（如果有）
    template = None
    if template_id:
        template = AIPromptTemplate.query.get(template_id)
    
    # 返回流式响应
    return ai_service.stream_chat_completion(ai_session, user_message_content, template, template_variables)

@app.route('/api/ai/templates')
@login_required
def ai_templates():
    """获取所有可用的prompt模板"""
    templates = AIPromptTemplate.query.filter_by(is_active=True).all()
    templates_data = []
    for template in templates:
        templates_data.append({
            'id': template.id,
            'name': template.name,
            'content': template.content,
            'description': template.description,
            'variables': template.variables or []
        })
    
    return jsonify(templates_data)

@app.route('/api/ai/models')
@login_required
def ai_models():
    """获取所有可用的模型（根据用户权限过滤）"""
    try:
        models = AIModelConfig.query.filter_by(is_active=True).all()
        models_data = []
        
        # 获取当前用户角色（处理未登录情况）
        if current_user.is_authenticated:
            user_role = current_user.role or 'user'  # 确保角色不为 None
            is_authenticated = True
        else:
            user_role = 'guest'
            is_authenticated = False
        
        for model in models:
            # 检查用户是否有权限使用该模型
            if model.can_access(user_role, is_authenticated):
                models_data.append({
                    'id': model.id,
                    'name': model.name,
                    'api_url': model.api_url,
                    'model_name': model.model_name,
                    'max_tokens': model.max_tokens,
                    'temperature': model.temperature,
                    'priority': model.priority
                })
        
        # 按优先级排序
        models_data.sort(key=lambda x: x.get('priority', 0), reverse=True)
        
        return jsonify(models_data)
    except Exception as e:
        print(f"Error in ai_models: {str(e)}")
        return jsonify({'error': '获取模型列表失败'}), 500

# 为未登录用户也提供可用的模型
@app.route('/api/ai/public-models')
def ai_public_models():
    """获取公开可用的模型（不需要登录）"""
    try:
        models = AIModelConfig.query.filter_by(is_active=True).all()
        models_data = []
        
        for model in models:
            # 检查模型是否允许未登录用户使用
            allowed_roles = model.allowed_roles or ['admin', 'teacher', 'student', 'user']
            if not model.require_login and 'guest' in allowed_roles:
                models_data.append({
                    'id': model.id,
                    'name': model.name,
                    'api_url': model.api_url,
                    'model_name': model.model_name,
                    'max_tokens': model.max_tokens,
                    'temperature': model.temperature,
                    'priority': model.priority
                })
        
        return jsonify(models_data)
    except Exception as e:
        print(f"Error in ai_public_models: {str(e)}")
        return jsonify({'error': '获取公开模型列表失败'}), 500

# API路由 - 模型配置（支持权限控制）
@app.route('/api/admin/ai/models', methods=['GET', 'POST', 'PUT', 'DELETE'])
@login_required
def admin_ai_models():
    """模型配置管理API"""
    if not current_user.is_admin:
        return jsonify({'error': '无权限'}), 403
    
    try:
        if request.method == 'GET':
            models = AIModelConfig.query.all()
            models_data = []
            for model in models:
                models_data.append(model.to_dict())
            return jsonify(models_data)
        
        elif request.method == 'POST':
            data = request.get_json()
            if not data:
                return jsonify({'error': '无效的请求数据'}), 400
                
            new_model = AIModelConfig(
                name=data.get('name', ''),
                api_url=data.get('api_url', ''),
                model_name=data.get('model_name', ''),
                max_tokens=data.get('max_tokens', 4000),
                temperature=data.get('temperature', 0.7),
                is_active=data.get('is_active', True),
                allowed_roles=data.get('allowed_roles', ['admin', 'teacher', 'student', 'user']),
                priority=data.get('priority', 0),
                require_login=data.get('require_login', True)
            )
            db.session.add(new_model)
            db.session.commit()
            return jsonify({'success': True, 'id': new_model.id})
        
        elif request.method == 'PUT':
            data = request.get_json()
            if not data:
                return jsonify({'error': '无效的请求数据'}), 400
                
            model_id = data.get('id')
            if not model_id:
                return jsonify({'error': '缺少模型ID'}), 400
                
            model = AIModelConfig.query.get_or_404(model_id)
            
            model.name = data.get('name', model.name)
            model.api_url = data.get('api_url', model.api_url)
            model.model_name = data.get('model_name', model.model_name)
            model.max_tokens = data.get('max_tokens', model.max_tokens)
            model.temperature = data.get('temperature', model.temperature)
            model.is_active = data.get('is_active', model.is_active)
            model.allowed_roles = data.get('allowed_roles', model.allowed_roles or ['admin', 'teacher', 'student', 'user'])
            model.priority = data.get('priority', model.priority)
            model.require_login = data.get('require_login', model.require_login)
            
            db.session.commit()
            return jsonify({'success': True})
        
        elif request.method == 'DELETE':
            model_id = request.args.get('id')
            if not model_id:
                return jsonify({'error': '缺少模型ID'}), 400
                
            model = AIModelConfig.query.get_or_404(model_id)
            db.session.delete(model)
            db.session.commit()
            return jsonify({'success': True})
            
    except Exception as e:
        print(f"Error in admin_ai_models: {str(e)}")
        db.session.rollback()
        return jsonify({'error': '操作失败'}), 500

@app.route('/api/ai/history')
@login_required
def ai_history():
    """查询对话历史"""
    keyword = request.args.get('keyword', '')
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    
    # 构建查询
    query = AISession.query.filter_by(user_id=current_user.id)
    
    if keyword:
        # 查询包含关键词的会话
        query = query.join(AIMessage).filter(
            db.or_(
                AISession.title.ilike(f'%{keyword}%'),
                AIMessage.content.ilike(f'%{keyword}%')
            )
        )
    
    sessions = query.order_by(AISession.updated_at.desc()).all()
    
    sessions_data = []
    for s in sessions:
        sessions_data.append({
            'id': s.id,
            'title': s.title,
            'model_used': s.model_used,
            'created_at': s.created_at.isoformat(),
            'updated_at': s.updated_at.isoformat(),
            'last_message': s.messages.order_by(AIMessage.created_at.desc()).first().content[:100] + '...' if s.messages.first() else ''
        })
    
    return jsonify(sessions_data)

@app.route('/api/ai/clear-context/<int:session_id>', methods=['POST'])
@login_required
def clear_context(session_id):
    """清空会话上下文"""
    ai_session = AISession.query.filter_by(id=session_id, user_id=current_user.id).first_or_404()
    
    # 删除所有消息
    ai_session.messages.delete()
    db.session.commit()
    
    return jsonify({'success': True})

#HTML项目管理------------------------------------------------------------------------------------
import uuid
# 路由定义
@app.route('/html_view/index')
def html_index():
    pages = HTMLPage.query.all()
    return render_template('html_view/index.html', pages=pages)

@app.route('/html_view/view/<int:page_id>')
def html_view_page(page_id):
    """查看页面（新标签页打开）"""
    page = HTMLPage.query.get_or_404(page_id)
    return render_template('/html_view/view.html', page=page)

@app.route('/html_view/raw/<int:page_id>')
def html_raw_page(page_id):
    """直接提供原始HTML文件并设置正确的内容类型"""
    page = HTMLPage.query.get_or_404(page_id)
    file_path = os.path.join(app.config['UPLOAD_HTML_FOLDER'], page.filename)
    
    # 检查文件是否存在
    if not os.path.exists(file_path):
        abort(404)
    
    # 发送文件并设置正确的内容类型
    return send_file(
        file_path,
        mimetype='text/html',
        as_attachment=False  # 确保不会触发下载
    )
    
@app.route('/html_view/admin')
@login_required
def html_admin():
    pages = HTMLPage.query.all()
    return render_template('/html_view/admin.html', pages=pages)

@app.route('/html_view/add', methods=['GET', 'POST'])
@login_required
def html_add_page():
    if request.method == 'POST':
        title = request.form['title']
        file = request.files['html_file']
        
        if file and file.filename.endswith('.html'):
            # 生成唯一文件名：UUID + 原始扩展名
            ext = os.path.splitext(file.filename)[1]
            filename = f"{uuid.uuid4().hex}{ext}"
            
            file.save(os.path.join(app.config['UPLOAD_HTML_FOLDER'], filename))
            
            # 创建缩略图（这里使用默认图片，实际可扩展为截图）
            thumbnail = 'default_thumb.png'
            
            new_page = HTMLPage(title=title, filename=filename, thumbnail=thumbnail)
            db.session.add(new_page)
            db.session.commit()
            return redirect(url_for('html_admin'))
    
    return render_template('/html_view/edit.html', action='添加')

@app.route('/html_view/edit/<int:page_id>', methods=['GET', 'POST'])
@login_required
def html_edit_page(page_id):
    page = HTMLPage.query.get_or_404(page_id)
    
    if request.method == 'POST':
        page.title = request.form['title']
        file = request.files['html_file']
        
        if file and file.filename.endswith('.html'):
            # 生成唯一文件名：UUID + 原始扩展名
            ext = os.path.splitext(file.filename)[1]
            filename = f"{uuid.uuid4().hex}{ext}"
            
            # 删除旧文件
            old_file = os.path.join(app.config['UPLOAD_HTML_FOLDER'], page.filename)
            if os.path.exists(old_file):
                os.remove(old_file)
            
            # 保存新文件
            filename = secure_filename(file.filename)
            file.save(os.path.join(app.config['UPLOAD_HTML_FOLDER'], filename))
            page.filename = filename
        
        db.session.commit()
        return redirect(url_for('html_admin'))
    
    return render_template('/html_view/edit.html', page=page, action='编辑')

@app.route('/html_view/delete/<int:page_id>')
@login_required
def html_delete_page(page_id):
    page = HTMLPage.query.get_or_404(page_id)
    
    # 删除文件
    file_path = os.path.join(app.config['UPLOAD_HTML_FOLDER'], page.filename)
    if os.path.exists(file_path):
        os.remove(file_path)
    
    db.session.delete(page)
    db.session.commit()
    return redirect(url_for('html_admin'))

@app.route('/html_view/uploads/<filename>')
def html_uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_HTML_FOLDER'], filename)

#成绩分析-----------------------------------------------------------------------------------------
#学生成绩
@app.route('/student/analysis/<int:student_id>')
@login_required
def student_analysis(student_id):
    # 权限检查
    #if not (current_user.is_admin() or current_user.is_teacher() or 
    #        (current_user.is_student() and current_user.id == student_id)):
    #    return render_template('403.html'), 403
    
    student = User.query.get_or_404(student_id)
    
    # 获取学生所有考试结果，按时间倒序排序
    all_results = ExamResult.query.filter_by(user_id=student_id)\
                                .order_by(ExamResult.submitted_at.desc())\
                                .all()
    
    # 处理同一考试的多次参加情况
    exam_data = {}
    for result in all_results:
        exam_id = result.exam_id
        # 如果还没有记录该考试，或者这次成绩更好（排名更靠前）
        if exam_id not in exam_data or result.score > exam_data[exam_id]['score']:
            exam = Exam.query.get(exam_id)
            logger.info(f"no exam: {exam_id},{exam}")
            class_id = student.class_id
            
            # 获取该班级所有学生在本次考试中的最佳成绩
            class_students = User.query.filter_by(class_id=class_id, role='student').all()
            student_results = []
            
            # 收集班级每个学生在该考试中的最佳成绩
            for class_student in class_students:
                student_results_query = ExamResult.query.filter_by(
                    user_id=class_student.id,
                    exam_id=exam_id
                ).order_by(ExamResult.score.desc()).first()
                
                if student_results_query:
                    student_results.append({
                        'user_id': class_student.id,
                        'score': student_results_query.score
                    })
            
            # 按分数降序排序
            student_results.sort(key=lambda x: x['score'], reverse=True)
            
            # 计算当前学生在该班级中的排名
            rank = None
            for i, r in enumerate(student_results):
                if r['user_id'] == student.id:
                    rank = i + 1
                    break
            if exam:
                exam_data[exam_id] = {
                    'exam_name': exam.name,
                    'score': result.score,
                    'total_score': len(exam.question_ids.split(',')) * exam.score_per_question,
                    'rank': rank,
                    'class_size': len(student_results),  # 实际参加考试的学生数
                    'date': result.submitted_at.strftime('%Y-%m-%d'),
                    'exam_id': exam.id
                }
    
    # 转换为列表并按时间倒序排序
    exam_data_list = sorted(exam_data.values(), key=lambda x: x['date'], reverse=True)
    
    # 知识点统计（处理以"|"分隔的知识点）
    knowledge_stats = {}
    for result in all_results:
        answers = json.loads(result.answers)
        for qid, answer in answers.items():
            question = Question.query.get(qid)
            if question and question.knowledge_point:
                # 分割知识点
                kps = question.knowledge_point.split('|')
                for kp in kps:
                    kp = kp.strip()
                    if kp:  # 确保知识点不为空
                        if kp not in knowledge_stats:
                            knowledge_stats[kp] = {'correct': 0, 'total': 0}
                        
                        knowledge_stats[kp]['total'] += 1
                        if answer == question.answer:
                            knowledge_stats[kp]['correct'] += 1
    
    # 转换为正确率
    knowledge_list = []
    for kp, stats in knowledge_stats.items():
        correct_rate = stats['correct'] / stats['total'] * 100 if stats['total'] > 0 else 0
        knowledge_list.append({
            'knowledge': kp,
            'correct_rate': round(correct_rate, 1),
            'correct': stats['correct'],
            'total': stats['total']
        })
    
    # 按正确率排序
    knowledge_list.sort(key=lambda x: x['correct_rate'], reverse=True)
    
    all_score = 0
    for item in exam_data_list:
        all_score += item["score"]
    
    return render_template('student_analysis.html', 
                          student=student,
                          exam_data=exam_data_list,
                          knowledge_list=knowledge_list,
                          average_score=0 if not exam_data_list else round(all_score/len(exam_data_list), 2)
        )
                          
# 班级成绩
@app.route('/class/analysis/<int:class_id>')
@login_required
@permissions_required('is_admin', 'is_teacher')
def class_analysis(class_id):
    
    class_ = Class.query.get_or_404(class_id)
    students = User.query.filter_by(class_id=class_id, role='student').all()
    
    # 获取班级所有考试ID
    exam_ids = [e.id for e in Exam.query.join(ExamClassAssociation).filter(
        ExamClassAssociation.class_id == class_id
    ).all()]
    
    # 计算每个学生的平均分和考试情况
    student_scores = []
    for student in students:
        # 获取学生所有考试记录
        all_results = ExamResult.query.filter_by(user_id=student.id)\
                                     .filter(ExamResult.exam_id.in_(exam_ids))\
                                     .all()
        
        # 按考试分组，只保留每个考试的最高分
        exam_best_scores = {}
        for result in all_results:
            exam_id = result.exam_id
            if exam_id not in exam_best_scores or result.score > exam_best_scores[exam_id]['score']:
                exam = Exam.query.get(exam_id)
                exam_best_scores[exam_id] = {
                    'score': result.score,
                    'total': len(exam.question_ids.split(',')) * exam.score_per_question,
                    'date': result.submitted_at
                }
        
        # 计算平均分
        best_scores = list(exam_best_scores.values())
        total_score = sum(r['score'] for r in best_scores)
        avg_score = total_score / len(best_scores) if best_scores else 0
        
        # 获取最近一次考试（按时间）
        latest_exam = None
        if best_scores:
            latest_exam = max(best_scores, key=lambda x: x['date'])
        
        student_scores.append({
            'id': student.id,
            'name': student.usernick or student.username,
            'avg_score': round(avg_score, 1),
            'exam_count': len(best_scores),  # 实际参加的不同考试数量
            'latest_score': latest_exam['score'] if latest_exam else '无',
            'best_scores': best_scores  # 存储每个考试的最佳成绩
        })
    
    # 按平均分排序
    student_scores.sort(key=lambda x: x['avg_score'], reverse=True)
    
    # 计算班级平均分
    class_avg = sum(s['avg_score'] for s in student_scores) / len(student_scores) if student_scores else 0
    
    return render_template('class_analysis.html', 
                          class_=class_, 
                          student_scores=student_scores,
                          class_avg=round(class_avg, 1))

#虫部落url展示-----------------------------------------------------------------------------------------
@app.route('/chongbuluo_url')
def chongbuluo_url():
    urls = chongbuluo_URL.query.all()
    log_request_info('chongbuluo_url', current_user, '访问虫部落链接')
    return render_template('chongbuluo_url.html', urls=urls)

#运行python代码-----------------------------------------------------------------------------------------
import multiprocessing, queue
import ast
import sys, signal
from copy import deepcopy
import traceback
# ---------------- 子进程执行函数 ----------------
def run_user_code(code, inputs, result_queue):
    try:
        tracer = CodeTracer(code, inputs=inputs)
        result_queue.put({
            "status": "success",
            "steps": tracer.steps,
            "output": tracer.output_text
        })
    except Exception as e:
        msg = "执行错误：" + str(e) + "\n" + traceback.format_exc()
        result_queue.put({"status": "error", "message": msg})

class CodeTracer:
    def __init__(self, code, inputs=None):
        self.steps = []
        self.allowed_types = (int, float, str, list, bool, type(None), dict)
        self.loop_vars = set()

        self.locals = {}
        self.output_text = ""

        # 输入队列（供 input() 使用）
        self._inputs = list(inputs) if inputs is not None else []
        self._input_index = 0

        # 安全 globals（注入安全 input/print）
        self.globals = self._create_safe_env()

        self.prev_line = None

        self._analyze_loop_vars(code)

        # 捕获 stdout
        self._stdout = io.StringIO()
        old_stdout = sys.stdout

        sys.settrace(self.global_tracer)
        try:
            sys.stdout = self._stdout
            code_obj = compile(code, filename="<user_code>", mode="exec")
            sys.settrace(self.global_tracer)
            exec(code_obj, self.globals, self.locals)
        except Exception as e:
            raise Exception("执行错误：" + str(e))
        finally:
            sys.settrace(None)
            sys.stdout = old_stdout
            try:
                self.output_text = self._stdout.getvalue()
            except Exception:
                self.output_text = ""

    def _analyze_loop_vars(self, code):
        class LoopVarFinder(ast.NodeVisitor):
            def __init__(self):
                self.vars = set()

            def visit_For(self, node):
                # 只收集 for i in ... 的 i（保持你原语义）
                if isinstance(node.target, ast.Name):
                    self.vars.add(node.target.id)
                self.generic_visit(node)

            def visit_While(self, node):
                self.generic_visit(node)

        tree = ast.parse(code)
        finder = LoopVarFinder()
        finder.visit(tree)
        self.loop_vars = finder.vars

    # 安全 input：从 inputs 数组按次取；没输入则抛异常（提示前端）
    def _safe_input(self, prompt=""):
        # prompt 会被 print 到 stdout（模拟 input 行为）
        if prompt is None:
            prompt = ""
        if prompt:
            print(prompt, end="")
        if self._input_index >= len(self._inputs):
            raise Exception("运行时请求输入，但前端未提供足够的输入行（input() 次数超过输入行数）。")
        s = self._inputs[self._input_index]
        self._input_index += 1
        # input() 返回字符串（保留原样，包含空串）
        return "" if s is None else str(s)

    def _create_safe_env(self):
        # 定制 print：默认即可，因为 stdout 已被捕获
        safe_builtins = {
            'range': range,
            'int': int,
            'str': str,
            'bool': bool,
            'list': list,
            'len': len,
            'abs': abs,
            'max': max,
            'min': min,
            'chr': chr,
            'ord': ord,
            'print': print,
            'input': self._safe_input,
        }
        return {'__builtins__': safe_builtins}

    def global_tracer(self, frame, event, arg):
        if frame.f_code.co_filename != "<user_code>":
            return self.global_tracer  # 或 return None（更彻底：不追踪该帧的子调用）
        if event == 'line':
            self.process_line(frame)
        return self.global_tracer

    def process_line(self, frame):
        current_line = frame.f_lineno
        if current_line == self.prev_line:
            return

        variables = {
            k: self._safe_copy(v)
            for k, v in self.locals.items()
            if isinstance(v, self.allowed_types)
        }
        pointers = {k: v for k, v in variables.items()
                    if k in self.loop_vars or k == '__loop_index'}

        self.steps.append({
            'line': current_line,
            'variables': deepcopy(variables),
            'pointers': pointers
        })
        self.prev_line = current_line

    def _safe_copy(self, obj):
        if isinstance(obj, list):
            return [self._safe_copy(x) for x in obj]
        elif isinstance(obj, dict):
            return {k: self._safe_copy(v) for k, v in obj.items()}
        return obj

@app.route("/run", methods=["POST"])
def run_code():
    data = request.get_json() or {}
    code = data.get("code", "")
    inputs = data.get("inputs", [])  # 前端传入：["line1", "line2", ...]

    result_queue = multiprocessing.Queue()
    p = multiprocessing.Process(target=run_user_code, args=(code, inputs, result_queue))
    p.start()

    p.join(3)
    if p.is_alive():
        p.terminate()
        return jsonify({"status": "error", "message": "代码运行超过 3 秒限制！请缩小数据规模或检查是否存在死循环。"})

    try:
        result = result_queue.get(timeout=1)
    except queue_mod.Empty:
        result = {"status": "error", "message": "子进程未返回任何信息，可能发生未知错误！"}
    except Exception as e:
        result = {"status": "error", "message": "未知错误：" + str(e)}
    return jsonify(result)


#用户GPT历史记录-----------------------------------------------------------------------------------------
@app.route('/user_chat_history')
@login_required
@permissions_required('is_admin')
def user_chat_history():
    page = request.args.get('page', 1, type=int)
    per_page = 50  # 每页显示10条
    keyword = request.args.get('keyword', '').strip()
    username = request.args.get('username', '').strip()
    
    # 构建查询
    query = User.query
    
    # 根据用户名过滤
    if username:
        query = query.filter(User.username.ilike(f'%{username}%'))
    
    users = query.all()
    
    # 收集所有聊天记录
    all_chats = []
    for user in users:
        if user.chat_history:
            chat_history = json.loads(user.chat_history) if user.chat_history else []
            for chat in chat_history:
                chat_item = {
                    'user_id': user.id,
                    'username': user.username,
                    'usernick': user.usernick,
                    'question': chat.get('question', ''),
                    'response': chat.get('response', ''),
                    'datetime': chat.get('datetime', '')  # 假设聊天记录中有时间戳
                }
                # 根据关键字过滤
                if keyword:
                    if (keyword.lower() in chat_item['question'].lower() or 
                        keyword.lower() in chat_item['response'].lower()):
                        all_chats.append(chat_item)
                else:
                    all_chats.append(chat_item)
    
    # 按照时间排序（假设聊天记录中有timestamp字段）
    all_chats.sort(key=lambda x: x.get('datetime', ''), reverse=True)
    
    # 分页
    total_chats = len(all_chats)
    start_idx = (page - 1) * per_page
    end_idx = start_idx + per_page
    paginated_chats = all_chats[start_idx:end_idx]
    
    # 分页对象
    pagination = {
        'page': page,
        'per_page': per_page,
        'total': total_chats,
        'pages': (total_chats + per_page - 1) // per_page,
        'has_prev': page > 1,
        'has_next': end_idx < total_chats,
        'prev_num': page - 1,
        'next_num': page + 1
    }
    
    return render_template('user_chat_history.html', 
                         chats=paginated_chats, 
                         pagination=pagination,
                         keyword=keyword,
                         username=username)

#推荐阅读------------------------------------------------------------------------------------------------
@app.route('/recommended_readings')
@login_required
@permissions_required('is_admin')
def recommended_readings_page():
    readings = RecommendedReading.query.order_by(RecommendedReading.date_added.desc()).all()
    return render_template('recommended_readings/manage.html', readings=readings)

@app.route('/add_recommended_reading', methods=['GET', 'POST'])
@login_required
@permissions_required('is_admin')
def add_recommended_reading():
    if request.method == 'POST':
        title = request.form['title']
        url = request.form['url']
        image_url = request.form['image_url']
        new_reading = RecommendedReading(title=title, url=url, image_url=image_url)
        db.session.add(new_reading)
        db.session.commit()
        return redirect(url_for('recommended_readings_page'))
    return render_template('recommended_readings/add.html')

@app.route('/edit_recommended_reading/<int:id>', methods=['GET', 'POST'])
@login_required
@permissions_required('is_admin')
def edit_recommended_reading(id):
    reading = RecommendedReading.query.get_or_404(id)
    if request.method == 'POST':
        reading.title = request.form['title']
        reading.url = request.form['url']
        reading.image_url = request.form['image_url']
        db.session.commit()
        return redirect(url_for('recommended_readings_page'))
    return render_template('recommended_readings/edit.html', reading=reading)

@app.route('/delete_recommended_reading/<int:id>', methods=['POST'])
@login_required
@permissions_required('is_admin')
def delete_recommended_reading(id):
    reading = RecommendedReading.query.get_or_404(id)
    db.session.delete(reading)
    db.session.commit()
    return redirect(url_for('recommended_readings_page'))

# 公告栏-------------------------------------------------------------------------------------------------
@app.route('/test')
@permissions_required('is_admin')
def test_page():
    return render_template('notice/test.html')

@app.route('/')
@app.route('/notice')
@login_required
def notice_page():
    notices = Notice.query.order_by(Notice.date_created.desc()).all()  # 获取所有公告，按创建时间降序
    recommended_readings = RecommendedReading.query.order_by(RecommendedReading.date_added.desc()).all()  # 获取所有推荐阅读，按添加时间降序
    return render_template('notice/notice.html', notices=notices, recommended_readings=recommended_readings, current_user=current_user)

@app.route('/add_notice', methods=['GET', 'POST'])
@login_required
@permissions_required('is_admin')
def add_notice_page():
    if request.method == 'POST':
        title = request.form['title']
        content = request.form['content']
        new_notice = Notice(title=title, content=content)
        db.session.add(new_notice)
        db.session.commit()
        return redirect(url_for('notice_page'))
    return render_template('notice/add_notice.html')

@app.route('/edit_notice/<int:id>', methods=['GET', 'POST'])
@login_required
@permissions_required('is_admin')
def edit_notice_page(id):
    notice = Notice.query.get_or_404(id)
    if request.method == 'POST':
        notice.title = request.form['title']
        notice.content = request.form['content']
        db.session.commit()
        return redirect(url_for('notice_page'))
    return render_template('notice/edit_notice.html', notice=notice)

@app.route('/delete_notice/<int:id>', methods=['POST'])
@login_required
@permissions_required('is_admin')
def delete_notice(id):
    notice = Notice.query.get_or_404(id)
    db.session.delete(notice)
    db.session.commit()
    return redirect(url_for('notice_page'))

@app.route('/notice/<int:id>')
def notice_detail_page(id):
    notice = Notice.query.get_or_404(id)
    return render_template('notice/notice_detail.html', notice=notice)

#学习中心(markdown)---------------------------------------------------------------------------------

# 定义Markdown文件的根目录
MARKDOWN_ROOT = os.path.join(app.root_path, 'static', 'markdown-viewer', 'md')

#保存markdown学习资料中的内容
@app.route('/save_markdown', methods=['POST'])
@login_required
@permissions_required('is_admin')  # 仅管理员可保存
def save_markdown():
    logger.info(f"Markdown root directory: {MARKDOWN_ROOT}")
    data = request.get_json()
    if not data:
        logger.warning("保存Markdown失败：无效的请求，没有收到JSON数据。")
        return jsonify({'success': False, 'message': '无效的请求'}), 400
    
    path = data.get('path')
    # 如果path以/static/markdown-viewer/md/开头，则去掉这部分前缀
    prefix = '/static/markdown-viewer/md/'
    if path.startswith(prefix):
        path = path[len(prefix):]

    content = data.get('content')

    if not path or content is None:
        logger.warning("保存Markdown失败：缺少参数。")
        return jsonify({'success': False, 'message': '缺少参数'}), 400

    # 安全检查，防止路径遍历
    normalized_path = os.path.normpath(path)
    if '..' in normalized_path.split(os.sep):
        logger.warning(f"保存Markdown失败：检测到路径遍历攻击。路径={path}")
        return jsonify({'success': False, 'message': '无效的文件路径'}), 400

    # 构建完整的文件路径
    full_path = os.path.join(MARKDOWN_ROOT, normalized_path + '.md')
    logger.info(f"尝试保存Markdown文件。路径={full_path}")

    # 确保目录存在
    os.makedirs(os.path.dirname(full_path), exist_ok=True)

    try:
        with open(full_path, 'w', encoding='utf-8') as f:
            f.write(content)
        logger.info(f"Markdown文件保存成功。路径={full_path}")
        return jsonify({'success': True})
    except Exception as e:
        logger.error(f"保存Markdown文件失败：{e}")
        return jsonify({'success': False, 'message': str(e)}), 500




#后台管理-------------------------------------------------------------------------------------------------
@app.route('/admin/index')
@login_required
@permissions_required('is_admin')
def admin_index():
    return render_template('admin/index.html')

@app.route('/admin/menu')
@login_required
@permissions_required('is_admin')
def admin_menu():
    return render_template('admin/menu.html')

#课堂学习模块-------------------------------------------------------------------------------------------------
# 新增课件
@app.route('/create_learning_material', methods=['GET', 'POST'])
@login_required
def create_learning_material():
    if request.method == 'POST':
        title = request.form['title']  # 获取标题
        content = request.form['content']
        class_ids = request.form.getlist('class_ids')  # 获取选择的班级IDs
        created_by = current_user.id  # 使用当前用户ID

        new_material = LearningMaterial(title=title, content=content, created_by=created_by)
        for class_id in class_ids:
            class_obj = Class.query.get(class_id)
            new_material.classes.append(class_obj)  # 将班级关联到课件

        db.session.add(new_material)
        db.session.commit()
        return redirect(url_for('list_learning_materials'))
    
    classes = Class.query.all()
    return render_template('admin/create_learning_material.html', classes=classes)

# 编辑课件
@app.route('/edit_learning_material/<int:id>', methods=['GET', 'POST'])
@login_required
def edit_learning_material(id):
    material = LearningMaterial.query.get_or_404(id)  # 获取要编辑的课件
    if request.method == 'POST':
        title = request.form['title']  # 获取标题
        content = request.form['content']
        class_ids = request.form.getlist('class_ids')  # 获取选择的班级IDs

        # 更新课件信息
        material.title = title
        material.content = content

        # 先清空当前课件的班级关联
        material.classes = []

        # 将新选中的班级关联到课件
        for class_id in class_ids:
            class_obj = Class.query.get(class_id)
            material.classes.append(class_obj)  # 重新添加班级关联

        db.session.commit()  # 提交修改
        return redirect(url_for('list_learning_materials'))  # 跳转到课件列表

    classes = Class.query.all()  # 获取所有班级
    # 获取当前课件已经关联的班级 IDs
    class_ids = [class_obj.id for class_obj in material.classes]
    return render_template('admin/edit_learning_material.html', material=material, classes=classes, class_ids=class_ids)

# 删除课件
@app.route('/delete_learning_material/<int:id>', methods=['GET'])
def delete_learning_material(id):
    material = LearningMaterial.query.get_or_404(id)
    
    # 手动删除关联表中的数据
    db.session.query(LearningMaterialClass).filter(LearningMaterialClass.material_id == id).delete()
    
    # 删除 LearningMaterial 数据
    db.session.delete(material)
    db.session.commit()
    
    return redirect(url_for('list_learning_materials'))


# 保存课件内容
@app.route('/api/save_learning_material', methods=['POST'])
def save_learning_material():
    content = request.form.get('content')
    title = request.form.get('title')
    learning_material = LearningMaterial(content=content, title=title)
    db.session.add(learning_material)
    db.session.commit()
    return jsonify({'message': '学习资料已保存'})

# 获取课件内容
@app.route('/api/get_learning_material', methods=['GET'])
def get_learning_material():
    # 获取最新的一份学习资料
    learning_material = LearningMaterial.query.order_by(LearningMaterial.created_at.desc()).first()
    if learning_material:
        return learning_material.content
    return jsonify({'error': '学习资料未找到'}), 404

# 学生签到----------------------------------------------------------------------------------------------------
# 学生签到路由
@app.route('/sign_in', methods=['POST'])
@login_required
def sign_in():
    try:
        # 获取当前日期和用户sign_in
        today = date.today()
        user = current_user

        # 检查当天是否已经签到
        sign_in_record = SignInRecord.query.filter_by(user_id=user.id, sign_in_date=today).first()
        if sign_in_record and sign_in_record.status == '已签到':
            return jsonify({'message': '今天已经签到'}), 400

        # 如果没有签到记录，则创建新的签到记录
        if not sign_in_record:
            sign_in_record = SignInRecord(user_id=user.id, sign_in_date=today, sign_in_time=datetime.now(), status='已签到')
            db.session.add(sign_in_record)
        else:
            sign_in_record.sign_in_time = datetime.now()
            sign_in_record.status = '已签到'

        db.session.commit()
        return jsonify({'message': '签到成功', 'sign_in_time': sign_in_record.sign_in_time.strftime('%Y-%m-%d %H:%M:%S')}), 200
    except Exception as e:
        # 在控制台输出错误
        print(f"签到时发生错误: {str(e)}")
        # 返回错误信息给前端
        return jsonify({'error': '服务器内部错误，请稍后再试。', 'details': str(e)}), 500

# 检查当天是否签到
# 检查签到状态时使用更高效的查询
@app.route('/check_sign_in_status', methods=['GET'])
@login_required
def check_sign_in_status():
    try:
        today = date.today()
        
        # 使用 exists() 查询，更高效
        exists = db.session.query(
            SignInRecord.query.filter_by(
                user_id=current_user.id, 
                sign_in_date=today,
                status='已签到'
            ).exists()
        ).scalar()
        
        if exists:
            # 如果需要时间信息，再查询具体记录
            record = SignInRecord.query.filter_by(
                user_id=current_user.id, 
                sign_in_date=today
            ).first()
            return jsonify({
                'status': '已签到', 
                'sign_in_time': record.sign_in_time.strftime('%Y-%m-%d %H:%M:%S')
            }), 200
        else:
            return jsonify({'status': '未签到'}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# 查看签到记录（按日期和班级）
@app.route('/admin/sign_in_records', methods=['GET'])
@login_required
@permissions_required('is_admin')
def get_sign_in_records():
    try:
        class_name = request.args.get('class_name')  # 班级名称
        sign_in_date = request.args.get('sign_in_date')  # 查询的日期

        # 如果没有指定日期，默认为今天
        if not sign_in_date:
            sign_in_date = date.today().strftime('%Y-%m-%d')

        # 查找指定班级
        user_class = Class.query.filter_by(name=class_name).first()
        if not user_class:
            return jsonify({'message': '班级不存在'}), 400

        # 查找班级中的所有学生
        students = User.query.filter_by(class_id=user_class.id).all()

        # 构建学生签到记录列表
        records = []
        for student in students:
            # 查询每个学生在指定日期的签到记录
            sign_in_record = SignInRecord.query.filter_by(user_id=student.id, sign_in_date=sign_in_date).first()
            if sign_in_record:
                # 如果有签到记录，使用记录中的签到状态
                records.append(sign_in_record.to_dict())
            else:
                # 如果没有签到记录，默认状态为“未签到”
                records.append({
                    'id': None,
                    'user_id': student.id,
                    'username': student.username,
                    'usernick': student.usernick,
                    'sign_in_date': sign_in_date,
                    'sign_in_time': None,
                    'status': '未签到'
                })

        return jsonify(records), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# 定时任务：每天12点重置签到状态
from apscheduler.schedulers.background import BackgroundScheduler

def reset_sign_in_status():
    try:
        # 重置所有用户的签到状态为 "未签到"
        today = date.today()
        sign_in_records = SignInRecord.query.filter_by(sign_in_date=today).all()
        for record in sign_in_records:
            record.status = '未签到'
        db.session.commit()
        print('签到状态重置成功')
    except Exception as e:
        print(f'重置签到状态时出错: {e}')

scheduler = BackgroundScheduler()
scheduler.add_job(func=reset_sign_in_status, trigger="cron", hour=0, minute=0)  # 每天午夜执行任务
scheduler.start()

# 博客功能----------------------------------------------------------------------------------------
from markupsafe import Markup
from markdown_it import MarkdownIt
from mdit_py_plugins.amsmath import amsmath_plugin

# 管理员能查看所有人的博客
@app.route('/admin/blogs', methods=['GET'])
@login_required
@permissions_required('is_admin')
def admin_blog_list():
    search_keyword = request.args.get('keyword', '')
    category = request.args.get('category', '')
    author_name = request.args.get('author', '')
    page = request.args.get('page', 1, type=int)
    per_page = 9
    
    #admin的博客只有自己能管理
    if current_user.id==1:
        query = Blog.query
    else:
        query = Blog.query.options(joinedload(Blog.author)).filter(
            or_(
                Blog.author_id != 1 
            )
        )
    if search_keyword:
        query = query.filter(Blog.title.contains(search_keyword) | 
                             Blog.content.contains(search_keyword) |
                             Blog.tags.contains(search_keyword))
    if category:
        query = query.filter_by(category=category)
    if author_name:
        query = query.join(User).filter(User.username.contains(author_name))

    blogs = query.paginate(page=page, per_page=per_page, error_out=False)
    log_request_info('admin/blogs', current_user, '查看博客')
    return render_template('blog/admin_blogs.html', blogs=blogs.items, pagination=blogs)

#管理员能删除所有人博客
@app.route('/admin/blogs/<int:blog_id>/delete', methods=['POST'])
@login_required
@permissions_required('is_admin')
def admin_delete_blog(blog_id):
    log_request_info('delete/blogs'+str(blog_id), current_user, '删除博客'+str(blog_id))
    blog = Blog.query.get_or_404(blog_id)
    db.session.delete(blog)
    db.session.commit()

    return jsonify({'message': '博客已删除'})

# 博客管理页面（后端）
@app.route('/blog/manage', methods=['GET'])
@login_required
def manage_blogs():
    log_request_info('blog/manage', current_user, '查看博客管理')
    page = request.args.get('page', 1, type=int)
    per_page = 9
    blogs = Blog.query.filter_by(author_id=current_user.id).paginate(page=page, per_page=per_page, error_out=False)
    return render_template('blog/manage.html', blogs=blogs.items, pagination=blogs)

# 新增博客页面
@app.route('/blog/new', methods=['GET', 'POST'])
@login_required
def new_blog():
    if request.method == 'POST':
        data = request.form
        title = data.get('title')
        content = data.get('content')
        category = data.get('category')
        tags = data.get('tags')
        is_public = 'is_public' in data

        blog = Blog(title=title, content=content, category=category, tags=tags, author_id=current_user.id, is_public=is_public)
        db.session.add(blog)
        db.session.commit()
        return redirect(url_for('manage_blogs'))

    return render_template('blog/new.html')


# 编辑博客
@app.route('/blog/edit/<int:blog_id>', methods=['GET', 'POST'])
@login_required
def edit_blog(blog_id):
    blog = Blog.query.get_or_404(blog_id)
    if blog.author_id != current_user.id:
        return render_template('403.html'), 403

    if request.method == 'POST':
        data = request.form
        blog.title = data.get('title')
        blog.content = data.get('content')
        blog.category = data.get('category')
        blog.tags = data.get('tags')
        blog.is_public = 'is_public' in data
        db.session.commit()
        return redirect(url_for('manage_blogs'))

    return render_template('blog/edit.html', blog=blog)


# 删除博客
@app.route('/blog/delete/<int:blog_id>', methods=['POST'])
@login_required
def delete_blog(blog_id):
    blog = Blog.query.get_or_404(blog_id)
    if blog.author_id != current_user.id:
        return render_template('403.html'), 403
    db.session.delete(blog)
    db.session.commit()
    return redirect(url_for('manage_blogs'))

# 前端博客列表页面
@app.route('/blogs', methods=['GET'])
@login_required
def blog_list():
    search_keyword = request.args.get('keyword', '')
    category = request.args.get('category', '')
    page = request.args.get('page', 1, type=int)
    per_page = 9
    
    #获取当前用户的或者其他用户已公开的博客
    query = Blog.query.options(joinedload(Blog.author)).filter(
        or_(
            Blog.is_public == True,
            Blog.author_id == current_user.id
        )
    )
    if search_keyword:
        query = query.filter(Blog.title.contains(search_keyword) | 
                             Blog.content.contains(search_keyword) |
                             Blog.tags.contains(search_keyword))
    if category:
        query = query.filter_by(category=category)

    blogs = query.paginate(page=page, per_page=per_page, error_out=False)
    log_request_info('blog/list', current_user, '查看博客列表')
    return render_template('blog/list.html', blogs=blogs.items, pagination=blogs)


# 博客详情页面
@app.route('/blog/<int:blog_id>', methods=['GET'])
@login_required
def blog_detail(blog_id):
    blog = Blog.query.options(joinedload(Blog.author)).get_or_404(blog_id)
    blog.view_count += 1
    db.session.commit()
    log_request_info('blog/detail'+str(blog_id), current_user, '查看博客文章'+str(blog_id))
    return render_template('blog/detail.html', blog=blog)

# 博客点赞
@app.route('/blog/like/<int:blog_id>', methods=['POST'])
@login_required
def like_blog(blog_id):
    blog = Blog.query.get_or_404(blog_id)
    blog.like_count += 1
    db.session.commit()
    #log_request_info('blog/like'+str(blog_id), current_user, '点赞博客'+str(blog_id))
    return jsonify({'message': '点赞成功', 'like_count': blog.like_count})


# 论坛功能----------------------------------------------------------------------------------------
# 获取帖子列表
@app.route('/post/', methods=['GET', 'POST'])
@login_required  # 需要登录才能访问
def get_post():
    if request.method == 'GET':
        # 如果是GET请求，渲染post.html页面
        return render_template('post.html')
    else:
        # 如果是POST请求，获取表单中的标题和内容，并创建新的帖子
        title = request.form.get('title')
        content = request.form.get('content')
        post = Post(title=title, content=content)
        user_id = current_user.id
        user = User.query.filter(User.id == user_id).first()
        post.author = user
        db.session.add(post)
        db.session.commit()
        # 重定向到论坛页面
        return redirect(url_for('forum_page'))

# 获取帖子详情
@app.route('/post_detail/<post_id>')
@login_required  # 需要登录才能访问
def post_detail_page(post_id):
    # 根据post_id查询帖子详情，并渲染post_detail.html页面
    post_model = Post.query.filter(Post.id == post_id).first()
    log_request_info('post_detail/'+str(post_id), current_user, '查看帖子'+str(post_id))
    return render_template('post_detail.html', post=post_model, post_count=len(post_model.replys))

# 新增回复
@app.route('/add_reply/', methods=['POST'])
@login_required  # 需要登录才能访问
def add_reply():
    # 获取表单中的回复内容和帖子ID，并创建新的回复
    content = request.form.get('reply_content')
    post_id = request.form.get('post_id')

    reply = Reply(content=content)
    user_id = current_user.id
    user = User.query.filter(User.id == user_id).first()
    reply.author = user
    post = Post.query.filter(Post.id == post_id).first()
    reply.post = post
    db.session.add(reply)
    db.session.commit()
    # 重定向到帖子详情页面
    return redirect(url_for('post_detail_page', post_id=post_id))

# 查询帖子
@app.route('/search_post/', methods=['GET'])
def search_post():
    # 获取搜索关键字，并根据标题、内容或作者用户名进行搜索
    search_key = request.args.get('q')
    context = {
        'posts': Post.query.filter(
            or_(
                Post.title.contains(search_key),
                Post.content.contains(search_key),
                Post.author.has(username=search_key),
                Post.author.has(usernick=search_key)
            )
        ).order_by(desc(Post.create_time)).all()
    }
    # 渲染论坛页面，显示搜索结果
    return render_template('forum.html', **context, len=len)

# 删除帖子
@app.route('/delete_post/<post_id>', methods=['POST'])
@login_required  # 需要登录才能访问
def delete_post(post_id):
    # 根据post_id查询帖子，如果帖子存在且当前用户有权限删除，执行删除操作
    post = Post.query.get_or_404(post_id)
    if current_user.role not in ['admin', 'teacher'] and post.author_id != current_user.id:
        return redirect(url_for('post_detail_page', post_id=post_id, message="无权限删除！"))
    
    db.session.delete(post)
    db.session.commit()
    # 重定向到论坛页面，并显示删除成功信息
    return redirect(url_for('forum_page', message="帖子已删除！"))

# 删除回复
@app.route('/delete_reply/<reply_id>', methods=['POST'])
@login_required  # 需要登录才能访问
def delete_reply(reply_id):
    # 根据reply_id查询回复，如果回复存在且当前用户有权限删除，执行删除操作
    reply = Reply.query.get_or_404(reply_id)
    if current_user.role not in ['admin', 'teacher'] and reply.author_id != current_user.id:
        return redirect(url_for('post_detail_page', post_id=reply.post_id, message="无权限删除！"))
    
    db.session.delete(reply)
    db.session.commit()
    # 重定向到帖子详情页面，并显示删除成功信息
    return redirect(url_for('post_detail_page', post_id=reply.post_id, message="回复已删除!"))

# 资料管理----------------------------------------------------------------------------------------
# 文件上传
@app.route('/upload_document', methods=['POST'])
@login_required  # 需要登录才能访问
@permissions_required('is_admin', 'is_teacher')  # 需要管理员或老师权限
def upload_document():
    # 获取节ID和上传的文件，并保存文件
    section_id = request.form['section_id']
    file = request.files['file']
    filename = file.filename
    file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))

    # 创建新的学习材料记录，并保存到数据库
    new_material = Material(name=filename, filename=filename, section_id=section_id)
    db.session.add(new_material)
    db.session.commit()
    # 重定向到资料页面
    return redirect(url_for('document_page'))

# 文件下载
@app.route('/download_document/<filename>')
@login_required  # 需要登录才能访问
def download_document(filename):
    # 从目录中发送文件进行下载
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)
    
# 删除文件
@app.route('/delete_document/<int:material_id>', methods=['POST'])
@login_required  # 需要登录才能访问
@permissions_required('is_admin', 'is_teacher')  # 需要管理员或老师权限
def delete_document(material_id):
    # 根据material_id查询学习材料记录，并删除文件和记录
    material = Material.query.get_or_404(material_id)
    try:
        os.remove(os.path.join(app.config['UPLOAD_FOLDER'], material.filename))
    except Exception as e:
        print(f"Error deleting file: {e}")

    db.session.delete(material)
    db.session.commit()
    # 重定向到资料页面
    return redirect(url_for('document_page'))

# 编辑文件信息
@app.route('/edit_document/<int:material_id>', methods=['POST'])
@login_required  # 需要登录才能访问
@permissions_required('is_admin', 'is_teacher')  # 需要管理员或老师权限
def edit_document(material_id):
    # 根据material_id查询学习材料记录，并更新名称和备注
    material = Material.query.get_or_404(material_id)
    new_name = request.form['new_name']
    new_note = request.form['new_note']

    material.name = new_name
    material.note = new_note
    db.session.commit()
    # 返回成功信息
    return jsonify(success=True)

# 获取文件信息
@app.route('/get_material/<int:material_id>', methods=['GET'])
@login_required  # 需要登录才能访问
def get_material(material_id):
    # 根据material_id查询学习材料记录，并返回其信息
    material = Material.query.get_or_404(material_id)
    return jsonify(name=material.name, note=material.note)

# 考试功能部分----------------------------------------------------------------------------------------
# 创建试卷
@app.route('/create_exam', methods=['POST'])
@login_required  # 需要登录才能访问
@permissions_required('is_admin', 'is_teacher')  # 需要管理员或老师权限
def create_exam():
    # 从请求中获取JSON数据，并创建新的Exam对象
    data = request.json
    classes = data.get('classes', [])

    exam = Exam(
        name=data['name'],
        question_ids=','.join(map(str, data['question_ids'])),
        duration=data['duration'],
        num_questions=data['num_questions'] if data['num_questions'] else 1,
        score_per_question=data['score_per_question'],
        creator_id=current_user.id,
        created_at=datetime.now()
    )

    db.session.add(exam)
    db.session.commit()

    # 为每个班级创建ExamClassAssociation对象并添加到数据库
    for class_name in classes:
        class_obj = Class.query.filter_by(name=class_name).first()
        if class_obj:
            association = ExamClassAssociation(exam_id=exam.id, class_id=class_obj.id)
            db.session.add(association)

    db.session.commit()
    # 返回试卷创建成功的信息
    return jsonify({'message': '试卷创建成功', 'exam_id': exam.id})

# 获取所有试卷
@app.route('/exams', methods=['GET'])
@login_required  # 需要登录才能访问
@permissions_required('is_admin', 'is_teacher', 'is_student')
def get_exams():
    try:
        # 根据用户角色获取相应的试卷，并按创建时间降序排序
        if current_user.role == 'admin':
            exams = Exam.query.order_by(Exam.created_at.desc()).all()
        elif current_user.role == 'teacher':
            exams = Exam.query.filter_by(creator_id=current_user.id).order_by(Exam.created_at.desc()).all()
        else:
            exams = Exam.query.join(ExamClassAssociation).join(Class).filter(
                ExamClassAssociation.class_id == current_user.class_id
            ).order_by(Exam.created_at.desc()).all()

        # 获取当前用户在所有试卷上的做题统计
        exam_ids = [exam.id for exam in exams]
        stats = db.session.query(
            ExamResult.exam_id,
            func.count(ExamResult.id).label('attempts'),
            func.max(ExamResult.score).label('max_score')
        ).filter(
            ExamResult.user_id == current_user.id,
            ExamResult.exam_id.in_(exam_ids)
        ).group_by(ExamResult.exam_id).all()

        # 将统计结果转换为字典，方便后续使用
        stats_dict = {stat.exam_id: {'attempts': stat.attempts, 'max_score': stat.max_score} for stat in stats}

        exams_list = []
        for exam in exams:
            exam_dict = exam.to_dict()
            classes = [[cls.id, cls.name] for cls in exam.classes]
            exam_dict['class'] = classes
            # 添加用户的做题统计
            user_stats = stats_dict.get(exam.id, {'attempts': 0, 'max_score': None})
            exam_dict['user_attempts'] = user_stats['attempts']
            exam_dict['user_max_score'] = user_stats['max_score']

            exams_list.append(exam_dict)

        return jsonify(exams_list)
    except Exception as e:
        # 返回错误信息
        return jsonify({'error': str(e)}), 500

# 返回单套试卷
@app.route('/api/exam/<int:exam_id>', methods=['GET'])
@login_required  # 需要登录才能访问
@permissions_required('is_admin', 'is_teacher')
def get_exam_data(exam_id):
    # 根据exam_id获取Exam对象，如果找不到则返回404错误
    exam = Exam.query.get_or_404(exam_id)
    if exam.creator_id != current_user.id and not current_user.is_admin:
        return render_template('403.html'), 403

    return jsonify(exam.to_dict())

# 单套试卷管理
@app.route('/manage_exam/<int:exam_id>', methods=['GET', 'POST'])
@login_required  # 需要登录才能访问
@permissions_required('is_admin', 'is_teacher')
def manage_exam(exam_id):
    # 根据exam_id获取Exam对象，如果找不到则返回404错误
    exam = Exam.query.get_or_404(exam_id)

    if exam.creator_id != current_user.id and not current_user.is_admin:
        return render_template('403.html'), 403

    if request.method == 'POST':
        try:
            data = request.json
            exam.name = data.get('name', exam.name)
            exam.question_ids = ','.join(map(str, data.get('question_ids', exam.question_ids.split(','))))
            exam.duration = data.get('duration', exam.duration)
            exam.num_questions = data.get('num_questions', exam.num_questions)
            exam.score_per_question = data.get('score_per_question', exam.score_per_question)
    
            # 处理班级更新
            classes = data.get('classes', [])
            exam.classes = []
            for class_id in classes:
                class_obj = Class.query.get(class_id)
                if class_obj:
                    exam.classes.append(class_obj)
    
            db.session.commit()
            return jsonify({'message': '更新成功！'})
    
        except Exception as e:
            db.session.rollback()
            return jsonify({'message': '服务器错误，请重试！'+str(e), 'error': str(e)}), 500


    return render_template('admin/exam_manage.html', exam=exam.to_dict())

# 删除试卷
@app.route('/delete_exam/<int:exam_id>', methods=['DELETE'])
@login_required  # 需要登录才能访问
@permissions_required('is_admin', 'is_teacher')
def delete_exam(exam_id):
    # 根据exam_id获取Exam对象，如果找不到则返回404错误
    exam = Exam.query.get_or_404(exam_id)
    if exam.creator_id != current_user.id and current_user.role!='admin':
        return render_template('403.html'), 403

    db.session.delete(exam)
    db.session.commit()
    # 返回试卷删除成功的信息
    return jsonify({'message': '试卷删除成功'})

# 用户点击“开始测试”后，生成试卷并显示题目
@app.route('/take_exam/<int:exam_id>', methods=['GET'])
@login_required  # 需要登录才能访问
def take_exam(exam_id):
    # 根据exam_id获取Exam对象，如果找不到则返回404错误
    exam = Exam.query.get_or_404(exam_id)
    
    # 非考试模式，直接允许多次尝试，不受考试时间和次数限制
    remaining_time = exam.duration * 60
    
    # 检查考试名称是否包含“考试”，并验证用户是否已经参加过该考试
    if ("测试" in exam.name or "考试" in exam.name) and current_user.role!='admin':
        existing_result = ExamResult.query.filter_by(user_id=current_user.id, exam_id=exam_id).first()
        if existing_result:
            flash("您已参加过此考试，不能重复参加。", "warning")
            return redirect(url_for('exam_list_page'))
        else:
            # 检查是否有未完成的考试会话
            exam_session = ExamSession.query.filter_by(user_id=current_user.id, exam_id=exam_id).first()
                
            if not exam_session:
                # 如果没有会话，则创建一个新的
                exam_session = ExamSession(user_id=current_user.id, exam_id=exam_id, start_time=datetime.now())
                db.session.add(exam_session)
                db.session.commit()
                
            # 计算剩余时间
            elapsed_time = (datetime.now() - exam_session.start_time).total_seconds()
            remaining_time = max(exam.duration * 60 - elapsed_time, 0)
                
            if remaining_time <= 0:
                flash("考试时间已到，无法重新进入。", "warning")
                return redirect(url_for('exam_list_page'))

    question_ids = list(map(int, exam.question_ids.split(',')))
    num_questions = min(exam.num_questions, len(question_ids))

    # 随机选择子集
    selected_question_ids = random.sample(question_ids, num_questions)

    # 根据原始顺序对选中的 ID 进行排序
    selected_question_ids_sorted = sorted(selected_question_ids, key=lambda x: question_ids.index(x))

    # 使用 SQLAlchemy 的 case 语句按照排序后的 ID 顺序查询
    order_case = case(
        {id: index for index, id in enumerate(selected_question_ids_sorted)},
        value=Question.id
    )

    questions = Question.query \
        .filter(Question.id.in_(selected_question_ids_sorted)) \
        .order_by(order_case) \
        .all()

    # 转换每个题目内容中的相对路径为绝对路径
    idx = 0
    for question in questions:
        question.content = convert_relative_paths_to_absolute(question.content)
        idx += 1

    log_request_info('take_exam/'+str(exam_id), current_user, '参加考试'+str(exam_id))
    return render_template('take_exam.html', exam=exam, questions=questions, remaining_time=remaining_time)

# 提交考试答案并显示结果
def check_answer(user_answer, answer, qtype):
    user_answer = user_answer or ""
    answer = answer or ""

    # 综合题（qtype=3）：去掉空格后整体匹配
    if qtype == 3:
        if user_answer.replace(' ', '') == answer.replace(' ', ''):
            return (1, 1)  # 完全正确
        else:
            return (0, 0)  # 不正确

    # 选择题：支持“多子题”（用空格分隔）
    user_answer = user_answer.split(" ")
    answer = answer.split(" ")

    n = len(user_answer)
    if user_answer == answer:
        return (1, n)

    c = 0
    for i in range(n):
        if i < len(answer) and user_answer[i] == answer[i]:
            c += 1

    if c:
        return (2, c)  # 部分正确（按子题计）
    else:
        return (0, 0)
        
from typing import Any, Dict, List, Tuple

def _sort_user_answers_by_exam_order(exam, user_answers: Dict[str, Any]) -> List[Tuple[int, str]]:
    """
    将用户答案按 Exam.question_ids 的顺序排序。
    返回 [(question_id:int, user_answer:str), ...]
    """
    question_ids = []
    if exam.question_ids:
        question_ids = list(map(int, exam.question_ids.split(',')))

    order_map = {qid: idx for idx, qid in enumerate(question_ids)}

    items: List[Tuple[int, str]] = []
    for k, v in (user_answers or {}).items():
        try:
            qid = int(k)
        except Exception:
            continue
        items.append((qid, v if v is not None else ""))

    items.sort(key=lambda kv: order_map.get(kv[0], 10**9))
    return items


def _build_results_payload(exam, sorted_items: List[Tuple[int, str]]):
    """
    根据 exam + 排好序的 (qid, user_answer) 生成 results 列表与 total_score（用于展示）。
    注意：这里不更新题目统计字段，不写数据库（结果页查看时使用）。
    """
    results = []
    total_score = 0

    for display_no, (question_id, user_answer) in enumerate(sorted_items, start=1):
        question = Question.query.get(question_id)
        if not question:
            results.append({
                'question_id': question_id,
                'display_no': display_no,
                'content': "<b>第" + str(display_no) + "题</b>（题目不存在：ID=" + str(question_id) + "）",
                'correct_answer': '',
                'user_answer': user_answer,
                'analysis': '',
                'is_correct': 0,
                'score': 0
            })
            continue

        correct_stat = check_answer(user_answer, question.answer, question.qtype)
        score = exam.score_per_question * correct_stat[1]
        total_score += score

        results.append({
            'question_id': question_id,
            'display_no': display_no,
            'content': "<b>第" + str(display_no) + "题</b>" + convert_relative_paths_to_absolute(question.content),
            'correct_answer': question.answer,
            'user_answer': user_answer,
            'analysis': convert_relative_paths_to_absolute(question.analysis),
            'is_correct': correct_stat[0],
            'score': score
        })

    return total_score, results

# 提交考试答案的路由
@app.route('/submit_exam', methods=['POST'])
@login_required
def submit_exam():
    try:
        data = request.get_json(silent=True) or {}
        exam_id = data.get('exam_id')
        user_answers = data.get('answers') or {}

        if not exam_id:
            return jsonify({'error': '缺少 exam_id'}), 400

        exam = Exam.query.get_or_404(exam_id)

        # 后端兜底：考试/测试模式禁止重复提交（避免绕过 take_exam）
        if ("测试" in exam.name or "考试" in exam.name) and current_user.role != 'admin':
            existing_result = ExamResult.query.filter_by(user_id=current_user.id, exam_id=exam_id).first()
            if existing_result:
                return jsonify({'error': '您已参加过此考试，不能重复提交。'}), 403

        sorted_items = _sort_user_answers_by_exam_order(exam, user_answers)

        total_score = 0
        # 提交时才更新题目统计字段
        for (question_id, user_answer) in sorted_items:
            question = Question.query.get(question_id)
            if not question:
                continue

            question.attempts = question.attempts or 0
            question.exam_attempts = question.exam_attempts or 0
            question.correct_answers = question.correct_answers or 0
            question.exam_correct_answers = question.exam_correct_answers or 0

            question.attempts += 1
            question.exam_attempts += 1

            correct_stat = check_answer(user_answer, question.answer, question.qtype)
            total_score += exam.score_per_question * correct_stat[1]

            if correct_stat[0] == 1:
                question.correct_answers += 1
                question.exam_correct_answers += 1

            db.session.add(question)

        exam_result = ExamResult(
            user_id=current_user.id,
            exam_id=exam_id,
            score=total_score,
            answers=json.dumps(user_answers, ensure_ascii=False)
        )
        db.session.add(exam_result)
        db.session.commit()

        return jsonify({
            'message': '考试提交成功',
            'redirect_url': url_for('view_exam_result', exam_result_id=exam_result.id)
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


# 提交考试后的结果查看页
@app.route('/exam_result/<int:exam_result_id>', methods=['GET'])
@login_required
def view_exam_result(exam_result_id):
    exam_result = ExamResult.query.get_or_404(exam_result_id)

    # 非管理员只能看自己的成绩
    if current_user.role != 'admin' and exam_result.user_id != current_user.id:
        return abort(403)

    exam = Exam.query.get_or_404(exam_result.exam_id)

    try:
        user_answers = json.loads(exam_result.answers or "{}")
    except Exception:
        user_answers = {}

    sorted_items = _sort_user_answers_by_exam_order(exam, user_answers)
    computed_total, results = _build_results_payload(exam, sorted_items)

    log_request_info('exam_result/' + str(exam_result_id), current_user, '查看考试结果' + str(exam_result_id))

    # total_score 优先用数据库里当时保存的 score（更“权威”）；computed_total 用于调试对照
    return render_template(
        'exam_result.html',
        exam=exam,
        exam_result=exam_result,
        total_score=exam_result.score,
        computed_total=computed_total,
        results=results
    )

# 返回当前用户的考试记录
@app.route('/my_exam_results/<int:exam_id>', methods=['GET'])
@login_required  # 需要登录才能访问
def my_exam_results(exam_id):
    # 根据exam_id和当前用户ID查询考试记录，并返回结果列表
    user_id = request.args.get('user_id', type=int)
    # 如果不带 user_id，就默认查看自己
    if not user_id:
        user_id = current_uer.id
        
    exam_results = ExamResult.query.filter_by(user_id=user_id, exam_id=exam_id).all()
    exam_name = Exam.query.filter_by(id=exam_id).first().name
    results = []
    for result in exam_results:
        results.append({
            'id': result.id,
            'exam_id': result.exam_id,
            'exam_name': exam_name,
            'score': result.score,
            'submitted_at': result.submitted_at.strftime('%Y-%m-%d %H:%M:%S')
        })
        
    return jsonify(results)

# 获取单个考试结果的详细信息
@app.route('/exam_detail/<int:result_id>', methods=['GET'])
@login_required  # 需要登录才能访问
@permissions_required('is_student', 'is_admin', 'is_teacher')
def exam_detail(result_id):
    user_id = request.args.get('user_id', type=int)
    # 如果不带 user_id，就默认查看自己
    if not user_id:
        user_id = current_uer.id
    exam_result = ExamResult.query.filter_by(id=result_id, user_id=user_id).first()
    
    if not exam_result:
        return jsonify({'error': 'Exam result not found'}), 404
    
    detailed_answers = exam_result.get_detailed_answers()
    return jsonify(detailed_answers)

import statistics
from sqlalchemy.orm import joinedload
# 获取试卷统计数据
@app.route('/exam_statistics/<int:exam_id>', methods=['GET'])
@login_required  # 需要登录才能访问
@permissions_required('is_admin', 'is_teacher', 'is_student')
def exam_statistics(exam_id):
    try:
        # 使用 joinedload 预加载关联的班级，减少后续的查询次数
        exam = Exam.query.options(joinedload(Exam.classes)).get(exam_id)
        if not exam:
            return jsonify({'message': '试卷不存在'}), 404

        question_ids = list(map(int, exam.question_ids.split(',')))
        ids_dict = {question_ids[i]: i + 1 for i in range(len(question_ids))}

        selected_class_id = request.args.get('class', '')
        if selected_class_id:
            # 前端筛选过的班级id
            exam_class_ids = [int(selected_class_id)]
        else:
            # 考试关联的班级ID列表
            exam_class_ids = [cls.id for cls in exam.classes]

        # 查询每个班级的总人数和已提交人数
        class_submissions = []
        for class_id in exam_class_ids:
            cls = Class.query.get(class_id)
            if not cls:
                continue
                
            # 获取班级总人数
            total_users = User.query.filter_by(class_id=class_id).count()
            
            # 获取班级已提交人数
            submitted_users = db.session.query(ExamResult.user_id)\
                .filter(ExamResult.exam_id == exam_id)\
                .join(User, User.id == ExamResult.user_id)\
                .filter(User.class_id == class_id)\
                .distinct()\
                .count()
            class_submissions.append({
                'class_id': class_id,
                'class_name': cls.name,
                'total': total_users,
                'submitted': submitted_users,
                'submission_rate': int((submitted_users / total_users * 100) if total_users > 0 else 0)
            })

        # 获取参加该考试的班级学生
        users = User.query.join(Class).filter(Class.id.in_(exam_class_ids)).all()

        # 优化 user_dict 的构建，使用字典推导
        user_dict = {
            user.id: [
                user.username,
                user.user_class.name if user.user_class else "未分班",
                user.usernick,
                user.email,
                user.user_class.id
            ]
            for user in users
        }

        user_ids = list(user_dict.keys())

        # 批量获取考试结果，避免多次查询
        results = ExamResult.query.filter_by(exam_id=exam.id)\
                                  .filter(ExamResult.user_id.in_(user_ids)).all()

        # 构建用户最高分字典
        user_max_score = {}
        for result in results:
            current_score = result.score
            if result.user_id not in user_max_score:
                user_max_score[result.user_id] = current_score
            else:
                if current_score > user_max_score[result.user_id]:
                    user_max_score[result.user_id] = current_score

        # 为未参加考试的用户设置分数为 '未参加'
        for user_id in user_ids:
            if user_id not in user_max_score:
                user_max_score[user_id] = '未参加'

        # 优化 user_scores_list 的构建，使用列表推导
        user_scores_list = [
            {
                'id': user_id,
                'username': user_dict[user_id][0],
                'class_name': user_dict[user_id][1],
                'usernick': user_dict[user_id][2],
                'highest_score': score,
                'email': user_dict[user_id][3],
                'class_id': user_dict[user_id][4],
                'exam_id': exam_id
            }
            for user_id, score in user_max_score.items()
        ]

        # 添加试卷总体统计信息
        scores = [score for score in user_max_score.values() if isinstance(score, int)]
        
        overall_stats = {
            #'highest_score': max(scores) if scores else 0,
            #'lowest_score': min(scores) if scores else 0,
            #'average_score': round(statistics.mean(scores), 2) if scores else 0,
            'score_distribution': {
                '0-29':   sum(1 for s in scores if s < 30),
                '30-39':  sum(1 for s in scores if 30 <= s < 40),
                '40-49':  sum(1 for s in scores if 40 <= s < 50),
                '50-59':  sum(1 for s in scores if 50 <= s < 60),
                '60-69':  sum(1 for s in scores if 60 <= s < 70),
                '70-79':  sum(1 for s in scores if 70 <= s < 80),
                '80-89':  sum(1 for s in scores if 80 <= s < 90),
                '90-100': sum(1 for s in scores if s >= 90)
            }
        }

        # 优化第1条：批量获取所有相关题目数据，避免在循环中逐个查询
        questions = Question.query.filter(Question.id.in_(question_ids)).all()
        questions_dict = {question.id: question for question in questions}

        # question_stats 仍先用 dict 来统计
        question_stats = {}
        for result in results:
            answers = json.loads(result.answers)
            for question_id_str, user_answer in answers.items():
                question_id = int(question_id_str)
                if question_id in question_ids:
                    question = questions_dict.get(question_id)
                    if not question:
                        # 题目已被删除或不存在
                        if question_id not in question_stats:
                            question_stats[question_id] = {
                                'content': f"【已删除的题目】ID: {question_id}",
                                'answer': "N/A",
                                'analysis': "",
                                'attempts': 0,
                                'correct_answers': 0,
                                'correct_rate': 0
                            }
                        continue
                    if question_id not in question_stats:
                        question_stats[question_id] = {
                            'content': question.content,
                            'answer': question.answer,
                            'analysis': question.analysis,
                            'attempts': 0,
                            'correct_answers': 0
                        }
                    question_stats[question_id]['attempts'] += 1
                    if question.answer == user_answer:
                        question_stats[question_id]['correct_answers'] += 1

        # 计算正确率、处理内容
        for qid, stats in question_stats.items():
            stats['correct_rate'] = round((stats['correct_answers'] / stats['attempts']) * 100, 2) if stats['attempts'] > 0 else 0
            stats['content'] = "<b>第" + str(ids_dict[qid]) + "题</b>" + convert_relative_paths_to_absolute(stats['content'])
            # 角色判断（省略），这里直接给解析
            stats['analysis'] = convert_relative_paths_to_absolute(stats['analysis'])

        # 按自定义规则对 question_stats 排序（得到一个 List[ (qid, stat), ... ]）
        sorted_question_stats = sorted(question_stats.items(), key=lambda item: ids_dict[item[0]])

        # 把排序后的数据重新塞回 question_stats（保持字典结构，不影响现有前端 $.each 用法）
        question_stats = dict(sorted_question_stats)

        # 额外做一个列表的形式（如果将来想前端严格按序显示，可直接用此列表）
        question_stats_list = [
            {
                'question_id': qid,
                **stat
            }
            for qid, stat in sorted_question_stats
        ]

        # 计算每个班级的统计指标
        class_statistics = []
        for class_id in exam_class_ids:
            # 获取该班级所有学生的分数
            class_scores = [score for user_id, score in user_max_score.items() 
                            if isinstance(score, int) and user_dict.get(user_id) and user_dict[user_id][4] == class_id]
            
            if class_scores:
                highest = max(class_scores)
                lowest = min(class_scores)
                average = round(statistics.mean(class_scores), 2)
                std_dev = round(statistics.stdev(class_scores), 2) if len(class_scores) > 1 else 0.0
            else:
                highest = lowest = average = std_dev = '无数据'
            
            class_statistics.append({
                'class_id': class_id,
                'highest': highest,
                'lowest': lowest,
                'average': average,
                'std_dev': std_dev
            })

        # 返回试卷统计数据，增加 question_stats_list
        return jsonify({
            'user_scores': user_scores_list,
            'question_stats': question_stats,              # 原字典结构（不影响旧代码）
            'question_stats_list': question_stats_list,    # 新增的列表结构
            'overall_stats': overall_stats,
            'class_submissions': class_submissions,  # 班级提交人数
            'class_statistics': class_statistics,
            'exam_id': exam_id
        })

    except Exception as e:
        logger.error(str(e))
        return jsonify({'error': f"发生异常: {str(e)}, 异常类型: {type(e)}, 角色: {current_user.role}"}), 500

#试卷统计中具体题目的学生做题情况
@app.route('/exam_statistics/<int:exam_id>/question/<int:question_id>/wrong_users', methods=['GET'])
@login_required
@permissions_required('is_admin', 'is_teacher')
def get_wrong_users(exam_id, question_id):
    try:
        # 获取试卷信息并预加载关联班级
        exam = Exam.query.options(joinedload(Exam.classes)).get(exam_id)
        if not exam:
            return jsonify({'message': '试卷不存在'}), 404

        selected_class_id = request.args.get('class', '')
        if selected_class_id:
            # 前端筛选过的班级id
            exam_class_ids = [int(selected_class_id)]
        else:
            # 考试关联的班级ID列表
            exam_class_ids = [cls.id for cls in exam.classes]
            
        users = User.query.join(Class).filter(Class.id.in_(exam_class_ids)).all()

        user_dict = {
            user.id: {
                'usernick': user.usernick
            }
            for user in users
        }

        user_ids = list(user_dict.keys())

        # 获取考试结果
        results = ExamResult.query.filter_by(exam_id=exam.id)\
                                  .filter(ExamResult.user_id.in_(user_ids)).all()

        # 获取题目答案
        question = Question.query.get(question_id)
        if not question:
            return jsonify({'message': '题目不存在'}), 404
        correct_answer = question.answer

        # 构建分组数据
        groups = {
            'A': [],
            'B': [],
            'C': [],
            'D': [],
            'Other': []
        }

        for result in results:
            answers = json.loads(result.answers)
            user_answer = answers.get(str(question_id), None)
            if user_answer is None:
                continue  # 可以选择是否包括未回答的用户

            group_key = user_answer if user_answer in ['A', 'B', 'C', 'D'] else 'Other'
            user_info = user_dict.get(result.user_id)
            if user_info:
                # 仅包含昵称和是否正确的标记
                group_entry = {
                    'user_id': result.user_id,
                    'usernick': user_info['usernick'],
                    'is_correct': user_answer == correct_answer
                }
                groups.setdefault(group_key, []).append(group_entry)

        return jsonify({'groups': groups, 'correct_answer': correct_answer}), 200

    except Exception as e:
        logger.error(str(e))
        return jsonify({'error': f"发生异常: {str(e)}, 异常类型: {type(e)}, 角色: {current_user.role}"}), 500




# 班级管理部分----------------------------------------------------------------------------------------
# 新增班级
@app.route('/add_classes', methods=['POST'])
@login_required  # 需要登录才能访问
@permissions_required('is_admin', 'is_teacher')
def add_class():
    try:
        # 获取请求数据并创建新的班级
        data = request.get_json()
        name = data.get('name')

        if not name:
            return jsonify({'error': 'Class name is required'}), 400

        existing_class = Class.query.filter_by(name=name).first()
        if existing_class:
            return jsonify({'error': 'Class name already exists'}), 400

        new_class = Class(name=name)
        db.session.add(new_class)
        db.session.commit()

        # 返回班级创建成功的信息
        return jsonify({'message': 'Class added successfully', 'class': {'id': new_class.id, 'name': new_class.name}}), 201
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# 获取所有班级
@app.route('/classes', methods=['GET'])
@login_required  # 需要登录才能访问
@permissions_required('is_admin', 'is_teacher', 'is_student')
def get_classes():
    classes = Class.query.all()
    classes_list = [{'id': cls.id, 'name': cls.name} for cls in classes]
    return jsonify(classes_list)

# 更新班级
@app.route('/classes/<int:class_id>', methods=['PUT'])
@login_required  # 需要登录才能访问
@permissions_required('is_admin', 'is_teacher')
def update_class(class_id):
    cls = Class.query.get_or_404(class_id)
    data = request.get_json()
    name = data.get('name')

    if not name:
        return jsonify({'error': 'Class name is required'}), 400

    existing_class = Class.query.filter_by(name=name).first()
    if existing_class and existing_class.id != class_id:
        return jsonify({'error': 'Class name already exists'}), 400

    cls.name = name
    db.session.commit()

    # 返回班级更新成功的信息
    return jsonify({'message': 'Class updated successfully', 'class': {'id': cls.id, 'name': cls.name}}), 200

# 删除班级
@app.route('/classes/<int:class_id>', methods=['DELETE'])
@login_required  # 需要登录才能访问
@permissions_required('is_admin', 'is_teacher')
def delete_class(class_id):
    cls = Class.query.get_or_404(class_id)
    db.session.delete(cls)
    db.session.commit()

    # 返回班级删除成功的信息
    return jsonify({'message': 'Class deleted successfully'}), 200

# 导入用户部分----------------------------------------------------------------------------------------
# 允许上传的文件扩展名
ALLOWED_EXTENSIONS = {'xlsx', 'docx'}

# 检查文件扩展名是否允许
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# 上传用户文件并解析
@app.route('/upload_users', methods=['GET', 'POST'])
@login_required  # 需要登录才能访问
@permissions_required('is_admin', 'is_teacher')
def upload_users():

    if request.method == 'POST':
        if 'file' not in request.files:
            return jsonify({"error": "No file part"})
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No selected file"})
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)

            if filename.rsplit('.', 1)[1].lower() == 'xlsx':
                df = pd.read_excel(file_path)
            elif filename.rsplit('.', 1)[1].lower() == 'docx':
                df = read_docx(file_path)

            for index, row in df.iterrows():
                class_name = row['class_name']
                class_obj = Class.query.filter_by(name=class_name).first()
                if not class_obj:
                    continue

                # 检查用户名是否已存在
                existing_user = User.query.filter_by(username=row['username']).first()
                if existing_user:
                    continue

                password = str(row['password'])

                user = User(
                    username=row['username'],
                    usernick=row['usernick'],
                    email=row['email'],
                    password=password,
                    class_id=class_obj.id,
                    role=row['role']
                )
                user.password = password
                db.session.add(user)
            db.session.commit()
            return jsonify({"message": "Users imported successfully"})
    return render_template('admin/upload_users.html')

# 解析docx文件
def read_docx(file_path):
    doc = Document(file_path)
    data = []
    keys = None
    for i, table in enumerate(doc.tables):
        if i == 0:
            keys = [cell.text for cell in table.rows[0].cells]
        for row in table.rows[1:]:
            data.append([cell.text for cell in row.cells])
    return pd.DataFrame(data, columns=keys)

# 项目文档页面的修改和保存
@app.route('/save-markdown', methods=['POST'])
@login_required  # 需要登录才能访问
@permissions_required('is_admin')
def save_help_markdown():
    markdown_content = request.form['markdown']
    with open('uploads/files/readme.md', 'w', encoding='utf-8') as f:
        f.write(markdown_content)
    return {'message': '保存成功！'}

# 匹配上传文件的网络路由
@app.route('/uploads/files/<filename>')
@login_required  # 需要登录才能访问
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

# 知识点API----------------------------------------------------------------------------------------
# 获取知识点列表
@app.route('/api/knowledgepoints')
@login_required  # 需要登录才能访问
def get_knowledge_points():
    # 查询所有模块，并按模块组织知识点
    modules = db.session.query(KnowledgePoint.module).distinct().all()

    data = []
    for module in modules:
        module_name = module[0]
        points = KnowledgePoint.query.filter_by(module=module_name).all()
        points_data = [point.to_dict() for point in points]
        data.append({'module': module_name, 'knowledgepoints': points_data})

    return jsonify(data)

# 添加知识点
@app.route('/api/add_knowledgepoint', methods=['POST'])
@login_required  # 需要登录才能访问
@permissions_required('is_admin', 'is_teacher')
def add_knowledge_point():
    if not request.json or not 'module' in request.json or not 'knowledgepoint' in request.json:
        return jsonify({'error': '数据格式错误'}), 400

    new_module = request.json['module']
    new_knowledgepoint = request.json['knowledgepoint']

    # 检查是否已存在相同的知识点
    existing_point = KnowledgePoint.query.filter_by(module=new_module, knowledgepoint=new_knowledgepoint).first()
    if existing_point:
        return jsonify({'error': '知识点已存在'}), 400

    # 创建并添加新的知识点
    kp = KnowledgePoint(module=new_module, knowledgepoint=new_knowledgepoint)
    db.session.add(kp)
    db.session.commit()

    return jsonify({'message': '知识点添加成功'}), 201

# 视频功能部分----------------------------------------------------------------------------------------
VIDEO_FOLDER = os.path.join(app.config['UPLOAD_FOLDER'], 'videos')
if not os.path.exists(VIDEO_FOLDER):
    os.makedirs(VIDEO_FOLDER)

# 视频上传API
@app.route('/upload_video', methods=['POST'])
@login_required  # 需要登录才能访问
@permissions_required('is_admin')
def upload_video():
    if 'file' not in request.files:
        return jsonify({'error': '无文件'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '未选择文件'}), 400
    if file:
        filename = secure_filename_with_chinese(file.filename)
        file_path = os.path.join(VIDEO_FOLDER, filename)
        file.save(file_path)
        return jsonify({'message': '上传成功'}), 200
    return jsonify({'error': '上传失败'}), 500

# 删除视频
@app.route('/delete_video/<filename>')
@login_required  # 需要登录才能访问
@permissions_required('is_admin')
def delete_video(filename):
    log_request_info('delete_video:' + filename,  current_user, '删除视频'+ filename)
    file_path = os.path.join(VIDEO_FOLDER, filename)
    if os.path.exists(file_path):
        os.remove(file_path)
        return jsonify({'message': '视频删除成功'}), 200
    else:
        return jsonify({'error': '视频未找到'}), 404

# 匹配视频上传路由
@app.route('/video/<filename>')
@login_required  # 需要登录才能访问
@permissions_required('is_admin')
def video(filename):
    return send_from_directory(VIDEO_FOLDER, filename)

# 工具网站API----------------------------------------------------------------------------------------
# 增加链接
@app.route('/api/add_link', methods=['POST'])
@login_required  # 需要登录才能访问
@permissions_required('is_admin')
def add_link():
    new_link = Link(
        url=request.form.get('url'),
        title=request.form.get('title'),
        description=request.form.get('description'),
        grp=request.form.get('group')
    )
    db.session.add(new_link)
    db.session.commit()
    return redirect(url_for('friend_page'))

# 更新链接
@app.route('/api/update_link/<int:link_id>', methods=['POST'])
@login_required  # 需要登录才能访问
@permissions_required('is_admin')
def update_link(link_id):
    data = request.json
    link = Link.query.get_or_404(link_id)
    link.url = data.get('url', link.url)
    link.title = data.get('title', link.title)
    link.description = data.get('description', link.description)
    db.session.commit()
    return jsonify({'message': '链接更新成功'}), 200

# 删除链接
@app.route('/api/delete_link/<int:link_id>')
@login_required  # 需要登录才能访问
@permissions_required('is_admin')
def delete_link(link_id):
    link = Link.query.get_or_404(link_id)
    db.session.delete(link)
    db.session.commit()
    return redirect(url_for('friend_page'))

# 问题管理部分----------------------------------------------------------------------------------------
import re
from sqlalchemy import text
from sqlalchemy import or_, and_
from sqlalchemy.sql import func

# 封装题目类型判定的函数
def determine_question_type(answer):
    if re.fullmatch(r'[A-D]', answer):
        return 'single_choice'
    elif re.fullmatch(r'([A-D] ?)+', answer):
        return 'multiple_choice'
    else:
        return 'composite'

#获取问题
@app.route('/api/questions', methods=['GET'])
@login_required  # 需要登录才能访问
def get_questions():
    try:
        # 获取分页参数
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 100, type=int)

        # 获取查询参数
        keyword = request.args.get('keyword', '', type=str)
        knowledge_point = request.args.get('knowledge_point', '', type=str)
        if '*' in knowledge_point:
            knowledge_point_type = 1  #1代表"与"关系
            knowledge_point = knowledge_point.split('*')
        else:
            knowledge_point_type = 0
            knowledge_point = knowledge_point.split('|')
        difficulty = request.args.get('difficulty', '', type=int)
        question_id = request.args.get('question_id', '', type=str)  # 批量ID作为字符串
        source = request.args.get('source', '', type=str).lower()
        exact_match = request.args.get('exact_match', 'false').lower() == 'true'
        choice_only = request.args.get('choice_only', 'false').lower() == 'true'
        choice_multi = request.args.get('choice_multi', 'false').lower() == 'true'
        fill_blank = request.args.get('fill_blank', 'false').lower() == 'true'
        my_questions = request.args.get('my_questions', 'false').lower() == 'true'

        # 构建查询条件
        query = Question.query
        
        # 如果 "我的题目" 选项被选择，过滤当前用户添加的问题
        if my_questions:
            query = query.join(Question.users).filter(User.id == current_user.id)

        # 批量ID查询
        if question_id:
            id_list = [int(qid) for qid in question_id.split(',') if qid.strip().isdigit()]  # 提取有效ID
                # 使用case语句按照原始顺序排序
            order_case = case({qid: index for index, qid in enumerate(id_list)}, value=Question.id)
            query = query.filter(Question.id.in_(id_list)).order_by(order_case)

        
        if keyword:
            keywords = keyword.split('|')
            if exact_match:
                query = query.filter(
                    and_(*[
                        or_(
                            Question.content.op('REGEXP')(f'{k}')
                        ) for k in keywords
                    ])
                )
            else:
                query = query.filter(
                    or_(*[
                        Question.content.ilike(f'%{k}%') |
                        Question.analysis.ilike(f'%{k}%') |
                        Question.answer.ilike(f'%{k}%') 
                        for k in keywords
                    ])
                )

        if knowledge_point:
            if knowledge_point_type:
                for kw in knowledge_point:
                    query = query.filter(Question.knowledge_point.contains(kw))
            else:
                filters = [Question.knowledge_point.contains(kw) for kw in knowledge_point]
                query = query.filter(or_(*filters))
        if difficulty:
            query = query.filter(Question.difficulty == difficulty)
        if source:
            query = query.filter(Question.source.ilike(f'%{source}%'))

        if choice_only or choice_multi or fill_blank:
            query = query.filter(
                or_(
                    and_(
                        #Question.answer.op('REGEXP')('^[A-D]$'),  # 单选题：单个字母A-D
                        Question.qtype==1,
                        choice_only
                    ),
                    and_(
                        #Question.answer.op('REGEXP')('^[A-D]( [A-D])+$'),  # 多选题：多个以空格分隔的字母A-D
                        Question.qtype==2,
                        choice_multi
                    ),
                    and_(
                        #~Question.answer.op('REGEXP')('^[A-D]( [A-D])*$'),  # 综合题：不匹配上述格式
                        or_(
                            Question.qtype==0,
                            Question.qtype==3
                        ),
                        fill_blank
                    )
                )
            )
        
        # 根据更新时间降序排序，如果时间相同则根据 ID 降序排序
        query = query.order_by(Question.updated_at.desc(), Question.id.desc())

        result_total = query.count()
        # 分页查询
        pagination = query.paginate(page=page, per_page=per_page, error_out=False)

        # 返回结果
        questions = pagination.items
        result = [{
            'id': q.id,
            'content': q.content,
            'knowledge_point': q.knowledge_point,
            'description': q.description,
            'difficulty': q.difficulty,
            'source': q.source,
            'answer': q.answer,
            'qtype': q.qtype,
            'analysis': q.analysis,
            'percentage': 0 if not q.attempts else round(q.correct_answers / q.attempts * 100, 1),
            'attempts': 0 if not q.attempts else q.attempts,
            'created_at':q.created_at,
            'updated_at':q.updated_at,
            'added_by': [user.username for user in q.users]  # 显示所有关联用户
        } for q in questions]

        return jsonify({
            'questions': result,
            'result_total': result_total,
            'total_pages': pagination.pages,
            'current_page': pagination.page,
            'has_next': pagination.has_next,
            'has_prev': pagination.has_prev
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# 删除问题
@app.route('/api/questions/<int:question_id>', methods=['DELETE'])
@login_required  # 需要登录才能访问
@permissions_required('is_admin')
def delete_question(question_id):
    question = Question.query.get(question_id)
    if not question:
        return jsonify({'error': '该问题不存在'}), 404

    db.session.delete(question)
    db.session.commit()
    return jsonify({'message': '删除成功'}), 200

# 批量删除问题
@app.route('/api/questions/batch_delete', methods=['DELETE'])
@login_required  # 需要登录才能访问
@permissions_required('is_admin')
def batch_delete_questions():
    try:
        question_ids = request.json.get('question_ids', [])
        if not question_ids:
            return jsonify({'error': 'No question IDs provided'}), 400
        
        # 删除题目
        Question.query.filter(Question.id.in_(question_ids)).delete(synchronize_session=False)
        db.session.commit()
        return jsonify({'message': f'Successfully deleted {len(question_ids)} questions.'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

# 查看具体某个问题
@app.route('/api/questions/<int:question_id>', methods=['GET'])
@login_required  # 需要登录才能访问
def get_question(question_id):
    question = Question.query.get(question_id)
    if not question:
        return jsonify({'error': '问题未找到'}), 404
    return jsonify(question.to_dict())

# 自动添加数据结构标签
def add_knowledge_point(knowledge_point):
    if ("栈" in knowledge_point or "队" in knowledge_point or "树" in knowledge_point or "链表" in knowledge_point) and ("数据结构" not in knowledge_point):
        return knowledge_point + "|数据结构"
    else:
        return knowledge_point

# 添加问题
@app.route('/api/questions', methods=['POST'])
@login_required  # 需要登录才能访问
@permissions_required('is_admin', 'is_teacher')
def add_question():
    try:
        content = request.form.get('content').replace('<span style="background-color: var(--bs-card-bg); color: var(--bs-modal-color); font-family: var(--bs-body-font-family); font-size: var(--bs-body-font-size); font-weight: var(--bs-body-font-weight); text-align: var(--bs-body-text-align);">', '')
        content_text = BeautifulSoup(content, 'html.parser').get_text().replace(" ", "").replace("\n", "")  # 解析出纯文本内容
        knowledge_point = add_knowledge_point(request.form.get('knowledge_point'))
        difficulty = request.form.get('difficulty')
        source = request.form.get('source')
        answer = request.form.get('answer')
        description = request.form.get('description')
        analysis = request.form.get('analysis').replace('<span style="background-color: var(--bs-card-bg); color: var(--bs-modal-color); font-family: var(--bs-body-font-family); font-size: var(--bs-body-font-size); font-weight: var(--bs-body-font-weight); text-align: var(--bs-body-text-align);">', '')
        qtype = request.form.get('qtype')
        if qtype == '0':  # 比较字符串类型
            answer = answer.strip()  # 去除前后空格
            if re.fullmatch(r'^[A-D]$', answer):  # 单选题
                qtype = '1'
            elif re.fullmatch(r'^[A-D]( [A-D])+$', answer):  # 多选题（至少两个选项）
                qtype = '2'
        if not content or not answer:
            return jsonify({'error': '题目和答案是必须的'}), 400
    
        # 检查重复或相似题目
        existing_questions = Question.query.all()
        content_bs4 = BeautifulSoup(content, 'html.parser').get_text()
        for existing_question in existing_questions:
            similarity = calculate_similarity(content_bs4, existing_question.content_text)
            if similarity > 0.8:
                return jsonify({
                    'error': '问题已存在',
                    'similar_questions': existing_question.id
                }), 401
    
        # 创建题目对象
        question = Question(
            content=content,
            content_text=content_text,  # 存储纯文本内容
            knowledge_point=knowledge_point,
            difficulty=difficulty,
            source=source,
            description=description,
            answer=answer,
            analysis=analysis,
            qtype=qtype,
            created_at=datetime.now(),
            updated_at=datetime.now()
        )
        
        # 添加题目到数据库
        db.session.add(question)
    
        # 将当前用户与该题目关联
        if question not in current_user.questions:
            current_user.questions.append(question)
        db.session.commit()
        
        return jsonify(question.to_dict()), 201
    except Exception as e:
        logger.error(str(e))
        return jsonify({"error": str(e)})

# 更新问题
@app.route('/api/questions/<int:question_id>', methods=['PUT', 'PATCH'])
@login_required  # 需要登录才能访问
@permissions_required('is_admin', 'is_teacher')
def edit_question(question_id):
    question = Question.query.get(question_id)
    if not question:
        return jsonify({'error': '不存在该问题'}), 404
    data = request.get_json()
    if not data:
        return jsonify({'error': '未获取到输入'}), 400

    # 关联用户与问题
    if current_user not in question.users:  # 避免用户重复记录
        current_user.questions.append(question)
        db.session.commit()  # 提交用户与问题的关联关系

    data['content'] = data['content'].replace('<span style="background-color: var(--bs-card-bg); color: var(--bs-modal-color); font-family: var(--bs-body-font-family); font-size: var(--bs-body-font-size); font-weight: var(--bs-body-font-weight); text-align: var(--bs-body-text-align);">', '')
    data['analysis'] = data['analysis'].replace('<span style="background-color: var(--bs-card-bg); color: var(--bs-modal-color); font-family: var(--bs-body-font-family); font-size: var(--bs-body-font-size); font-weight: var(--bs-body-font-weight); text-align: var(--bs-body-text-align);">', '')
    
    # 检查是否更新了`content`
    if 'content' in data and data['content'] != question.content:
        question.content = data['content']
        # 重新解析并更新`content_text`
        question.content_text = BeautifulSoup(data['content'], 'html.parser').get_text()

    # 更新其他字段
    question.update(data)
    question.updated_at = datetime.now()  # 更新问题时自动更新时间

    db.session.commit()
    return jsonify(question.to_dict()), 200


#问题查重
@app.route('/check_duplicates', methods=['GET'])
def check_duplicates():
    try:
        questions = Question.query.all()
        duplicates = []
        checked = set()  # 用于存储已检查的题目
    
        for question in questions:
            for other in questions:
                if question.id != other.id and (question.id, other.id) not in checked and question.content_text[0]==other.content_text[0] and abs(len(question.content_text)-len(other.content_text))<=10:
                    # 计算两题的相似度
                    similarity = calculate_similarity(question.content_text, other.content_text)
                    if similarity > 0.6:  # 假设相似度大于0.8为重复
                        duplicates.append({
                            "id1": question.id,
                            "id2": other.id,
                            "similarity": similarity
                        })
                    # 标记为已检查
                    checked.add((question.id, other.id))
                    checked.add((other.id, question.id))

        return jsonify({"duplicates": duplicates})
    except Exception as e:
        logger.error(str(e))
        return jsonify({"error": str(e)})
# 题目导入----------------------------------------------------------------------------------------
import os
import re
import pandas as pd
from docx import Document
import PyPDF2
import tempfile
import json
from io import BytesIO
from flask import jsonify, request, render_template
from werkzeug.utils import secure_filename
import traceback

EXTRACT_PROMPT = """
你是一个专业的试题解析AI，请根据提供的试卷内容，提取其中的所有试题，并按照以下格式输出：

### 要求：
1. 提取所有信息技术试题，包括选择题、材料选择题和综合题
2. 每道试题都用```html和```包裹，注意部分字符需要转义符表示，详见下表：
| 字符 | 转义符   |
| ---- | -------- |
| `<`  | `&lt;`   |
| `>`  | `&gt;`   |
| `&`  | `&amp;`  |
| `"`  | `&quot;` |
| `'`  | `&#39;`  |
| ` `(空格)  | `&nbsp;`  |
3. 每道试题包含以下字段（每行一个字段）：
   来源：[试卷名称或来源]
   题目：[题目内容（删除原有的题号），用HTML格式，保留换行和基本结构。题目中可能出现python代码，由于读取文档是大概率导致python的缩进丢失，所以请帮我整理格式，尽可能还原代码中4空格的缩进]
   题型：[1-单选题, 2-多选题, 3-综合题]
   答案：[选择题用大写字母，如A/AB/ABC；综合题为文本答案]
   解析：[如果有解析就填写，没有就留空]
   知识点：[试题涉及的知识点，包括：数据与信息 数据管理 数据安全 进制与编码 流程图 大数据 人工智能 pandas python基础 列表 字符串 字典 函数 循环 选择 枚举与解析 文本数据处理 信息系统 硬件与软件 传感与控制 网络 信息系统搭建 网络应用软件 数据结构 python大题 栈 队列 链表 二叉树 递归 排序 随机数 模拟 对分查找 时间复杂度 常规题 创新题 较差 一般 较好 极好 超纲 学考题 基础题。多个知识点之间用|隔开]
   难度：[1-5之间的整数，1最简单，5最难。1级：单纯的知识性考察  2级：日常的应用，初步的理解，简单的逻辑  3级：稍复杂的逻辑，原理性的理解  4级：对解决生活中一般问题能力的考察，或较强逻辑推演能力的考察 5级：对解决生活中复杂问题能力的考察，考察形式多样化，需极强的信息意识与逻辑能力，一般为python压轴大题。]

### 特别说明：
1. 题目格式要求：
   - 使用HTML格式，但不要添加任何样式
   - 对于python代码尽可能根据题目逻辑添加4空格的缩进
   - 换行可用<br>标签
   - 表格尽量用table、tr、td标签还原结构
   - 不保留图片，用[图片]标注
   - 单选、多选题、材料题的选项空采用（▲）标记

2. 材料选择题处理：
   - 如果是材料题（一段材料后有多道选择题），请将每道选择题都单独作为一道题
   - 每道选择题的题目都要包含相同的材料内容
   - 格式：先放材料，再放题目

3. 题型判断：
   - 单选题：只有一个正确答案的选择题，答案如A、B、C、D
   - 多选题：有多个正确答案的选择题，答案如AB、ACD，空格隔开
   - 综合题：包含解答、计算、论述等需要较长回答的题目

4. 输出格式示例：
   ```html
   来源：2023年2月强基联盟
   题目：某 Python 代码段如下，执行后输出的结果是（▲）<br>list1 =[2, 3, 4, 5, 6]<br>c1=1;c2=1<br>for n in list1:<br>&nbsp;&nbsp;&nbsp;&nbsp;if n % 3 == 0<br>&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;c1 *= n<br>&nbsp;&nbsp;&nbsp;&nbsp;else:<br>&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;c2 *= n<br>print(c1-c2)<br>A.-5<br>B.-8<br>C.-22<br>D.-358
   题型：1
   答案：C
   解析：给定的 Python 代码段会对列表 `list1` 进行遍历，对可以被 3 整除的元素进行乘积计算，最终输出 `c1` 和 `c2` 的差值。
   知识点：python基础|循环
   难度：2

请开始解析以下试卷内容：
{content}
"""

# 支持的文件类型
ALLOWED_EXTENSIONS = {'xlsx', 'docx', 'pdf'}

def allowed_file(filename):
    """检查文件类型是否允许"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def read_docx(filepath):
    """读取docx文件内容"""
    try:
        doc = Document(filepath)
        content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)
        
        # 读取表格
        for table in doc.tables:
            table_content = []
            for row in table.rows:
                row_content = []
                for cell in row.cells:
                    row_content.append(cell.text.strip())
                table_content.append(" | ".join(row_content))
            if table_content:
                content.append("[表格开始]")
                content.extend(table_content)
                content.append("[表格结束]")
        
        return "\n".join(content)
    except Exception as e:
        raise Exception(f"读取docx文件失败: {str(e)}")

def read_pdf(filepath):
    """读取pdf文件内容"""
    try:
        content = []
        with open(filepath, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text.strip():
                    content.append(f"=== 第{page_num + 1}页 ===")
                    content.append(text)
        return "\n".join(content)
    except Exception as e:
        raise Exception(f"读取pdf文件失败: {str(e)}")

def extract_questions_from_ai(content, model="gpt-5.1"):
    """使用AI提取试题"""
    try:
        # 构建完整的prompt
        prompt = EXTRACT_PROMPT.format(content=content)
        
        # 调用AI
        response = chatgpt(prompt, model)
        
        # 使用正则表达式提取所有试题块
        question_pattern = r'```html\s*(.*?)\s*```'
        question_blocks = re.findall(question_pattern, response, re.DOTALL)
        
        if not question_blocks:
            # 如果没有找到试题块，尝试直接解析响应
            return parse_ai_response_directly(response)
        
        questions = []
        for block in question_blocks:
            question_data = parse_question_block(block)
            if question_data:
                questions.append(question_data)
        
        return questions
    
    except Exception as e:
        raise Exception(f"AI解析失败: {str(e)}")

def parse_question_block(block):
    """解析单个试题块"""
    try:
        lines = block.strip().split('\n')
        question_data = {
            '来源': '未知',
            '题目': '',
            '题型': 3,  # 默认为综合题
            '答案': '',
            '解析': '',
            '知识点': '',
            '难度': 3
        }
        
        current_field = None
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # 检查是否是字段行
            field_match = re.match(r'^(来源|题目|题型|答案|解析|知识点|难度)：\s*(.*)', line)
            if field_match:
                field_name = field_match.group(1)
                field_value = field_match.group(2)
                question_data[field_name] = field_value
                current_field = field_name
            elif current_field and current_field in question_data:
                # 如果是多行内容，追加到当前字段
                question_data[current_field] += "<br>" + line
        
        # 转换题型为数字
        if isinstance(question_data['题型'], str):
            if '单选' in question_data['题型'] or question_data['题型'] == '1':
                question_data['题型'] = 1
            elif '多选' in question_data['题型'] or question_data['题型'] == '2':
                question_data['题型'] = 2
            elif '综合' in question_data['题型'] or question_data['题型'] == '3':
                question_data['题型'] = 3
            else:
                try:
                    question_data['题型'] = int(question_data['题型'])
                except:
                    question_data['题型'] = 3
        
        # 确保难度是数字
        if isinstance(question_data['难度'], str):
            try:
                question_data['难度'] = int(question_data['难度'])
            except:
                question_data['难度'] = 3
        
        return question_data
        
    except Exception as e:
        print(f"解析试题块失败: {str(e)}")
        return None

def parse_ai_response_directly(response):
    """直接解析AI响应（当没有找到```html块时）"""
    questions = []
    lines = response.strip().split('\n')
    
    current_question = None
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # 尝试识别新的题目开始
        if re.match(r'^题目\d*[：:]', line) or re.match(r'^\d+[\.、]', line):
            if current_question:
                questions.append(current_question)
            current_question = {
                '来源': 'AI解析',
                '题目': '',
                '题型': 3,
                '答案': '',
                '解析': '',
                '知识点': '',
                '难度': 3
            }
            current_question['题目'] = line
        
        elif current_question:
            # 识别答案行
            if re.match(r'^[答案][：:]', line) or '正确答案' in line:
                current_question['答案'] = line
            # 识别解析行
            elif re.match(r'^[解析][：:]', line) or '解析' in line:
                current_question['解析'] = line
            else:
                # 追加到题目
                current_question['题目'] += "<br>" + line
    
    if current_question:
        questions.append(current_question)
    
    return questions

def save_to_excel(questions, output_path):
    """将试题保存为Excel文件"""
    df = pd.DataFrame(questions)
    df.to_excel(output_path, index=False)
    return output_path

@app.route('/upload_questions', methods=['GET', 'POST'])
@login_required
@permissions_required('is_admin', 'is_teacher')
def upload_questions():
    if request.method == 'POST':
        try:
            # 检查文件是否存在
            if 'file' not in request.files:
                return jsonify({'error': '没有选择文件'}), 400
            
            file = request.files['file']
            if file.filename == '':
                return jsonify({'error': '没有选择文件'}), 400
            
            # 检查文件类型
            if not allowed_file(file.filename):
                return jsonify({'error': '不支持的文件类型，请上传xlsx、docx或pdf文件'}), 400
            
            # 检查文件大小（仅对pdf）
            if file.filename.endswith('.pdf'):
                file.seek(0, 2)  # 移动到文件末尾
                file_size = file.tell()
                file.seek(0)  # 重置文件指针
                if file_size > 8 * 1024 * 1024:  # 8MB
                    return jsonify({'error': 'PDF文件大小不能超过8MB'}), 400
            
            # 保存上传的文件
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            # 检查是否有解析文件
            analysis_content = None
            if 'analysis_file' in request.files:
                analysis_file = request.files['analysis_file']
                if analysis_file.filename != '' and allowed_file(analysis_file.filename):
                    analysis_filename = secure_filename(analysis_file.filename)
                    analysis_filepath = os.path.join(app.config['UPLOAD_FOLDER'], analysis_filename)
                    analysis_file.save(analysis_filepath)
                    
                    # 读取解析文件内容
                    if analysis_filename.endswith('.docx'):
                        analysis_content = read_docx(analysis_filepath)
                    elif analysis_filename.endswith('.pdf'):
                        analysis_content = read_pdf(analysis_filepath)
            
            # 处理不同的文件类型
            if filename.endswith('.xlsx'):
                # 直接导入xlsx
                wrong, succeed = import_questions(filepath)
                message = f'导入成功{succeed}条\n失败：{len(wrong)}条'
                if wrong:
                    message += '\n' + '\n'.join(wrong[:10])  # 只显示前10条错误
                    if len(wrong) > 10:
                        message += f'\n...还有{len(wrong)-10}条错误'
                return jsonify({'message': message}), 200
            
            else:
                # 读取文件内容
                content = ''
                if filename.endswith('.docx'):
                    content = read_docx(filepath)
                elif filename.endswith('.pdf'):
                    content = read_pdf(filepath)
                
                # 如果有解析内容，合并到试题内容中
                if analysis_content:
                    content += "\n\n=== 答案解析 ===\n" + analysis_content
                
                # 使用AI提取试题
                questions = extract_questions_from_ai(content)
                
                if not questions:
                    return jsonify({'error': '未能从文件中提取到试题'}), 400
                
                # 创建临时Excel文件
                temp_excel = tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False)
                excel_path = temp_excel.name
                temp_excel.close()
                
                # 保存为Excel
                save_to_excel(questions, excel_path)
                # 导入试题
                wrong, succeed = import_questions(excel_path)
                
                # 清理临时文件
                try:
                    os.unlink(excel_path)
                except:
                    pass
                
                # 返回结果
                message = f'AI解析成功，导入成功{succeed}条\n失败：{len(wrong)}条'
                if wrong:
                    message += '\n' + '\n'.join(wrong[:10])
                    if len(wrong) > 10:
                        message += f'\n...还有{len(wrong)-10}条错误'
                
                return jsonify({
                    'message': message,
                    'parsed_count': len(questions),
                    'imported_count': succeed,
                    'failed_count': len(wrong)
                }), 200
        
        except Exception as e:
            error_detail = traceback.format_exc()
            print(f"上传处理失败: {error_detail}")
            return jsonify({'error': f'处理失败: {str(e)}'}), 500
    
    return render_template('admin/upload_questions.html')

# 导入题目到数据库
def import_questions(filepath):
    data = pd.read_excel(filepath)
    wrong = []
    succeed = 0
    for index, row in data.iterrows():
        content=row['题目']
        answer = row['答案']
        qtype = row['题型']
        knowledge_point= row['知识点']
        #自动识别题型
        if qtype == 0:  # 比较字符串类型
            if re.fullmatch(r'^[A-D]$', answer.strip()):  # 单选题
                qtype = 1
            elif re.fullmatch(r'^[A-D]( [A-D])+$', answer.strip()):  # 多选题（至少两个选项）
                qtype = 2
            else:
                qtype = 3
        if not pd.isna(content) and not pd.isna(knowledge_point):
            succeed += 1
            question = Question(
                source=row['来源'],
                content=content,
                content_text=content, #BeautifulSoup(content, 'html.parser').get_text().replace(" ", "").replace("\n", ""),  # 解析出纯文本内容
                qtype=qtype,
                answer= answer,
                analysis= row['解析'],
                description='',
                knowledge_point=str(knowledge_point),
                difficulty=row['难度']
            )
            db.session.add(question)
        else:
            wrong.append("第" + str(index+1) + "条：" +('缺少题目' if pd.isna(content) else '') + ('缺少知识点' if pd.isna(knowledge_point) else ''))
    db.session.commit()
    return wrong, succeed

# HTML格式整理----------------------------------------------------------------------------------------
def clean(content):
        
    # 替换所有的全角小数点为半角
    content = content.replace('．', '.')  # 直接替换整个HTML内容中的全角小数点
    # 解析HTML内容
    soup = BeautifulSoup(content, 'lxml')
    
    # 清除所有<p>标签的样式
    for p_tag in soup.find_all(['p']):
        p_tag.name = 'div'  # 将标签名改为<div>
        p_tag.attrs = {}

    # 清除<span>标签中除了特定样式（border、font-family等）外的所有样式
    for span_tag in soup.find_all(['div', 'u', 'span']):
        style = span_tag.attrs.get('style', '')
        allowed_styles = []
        
        # 提取允许保留的样式：border, font-family, font-size
        for rule in style.split(';'):
            if any(key in rule for key in ['border:', 'font-family:']):
                allowed_styles.append(rule.strip())
        # 如果存在允许的样式，则保留这些样式，否则清空样式
        if allowed_styles:
            span_tag.attrs = {'style': '; '.join(allowed_styles)}
        else:
            span_tag.attrs = {}

    # 移除所有注释
    comments = soup.find_all(string=lambda text: isinstance(text, Comment))
    for comment in comments:
        comment.extract()

    # 移除无用的<span>标签，只保留内容
    for span_tag in soup.find_all('span'):
        if span_tag.attrs == {}:
            span_tag.unwrap()
    # 移除<o:p>标签，只保留内容
    for op_tag in soup.find_all('o:p'):
        op_tag.unwrap()
        
    # 将连续的<div><br></div>缩减为一个<br>
    for div_tag in soup.find_all('div'):
        # 如果<div>标签中只包含一个<br>标签，移除<div>标签，仅保留<br>
        if len(div_tag.contents) == 1 and div_tag.contents[0].name == 'br':
            div_tag.unwrap()
    # 再次清理<br>标签
    for br_tag in soup.find_all('br'):
        br_tag.decompose()  # 完全删除标签及其内容
    return str(soup)
    

# 清理HTML标签
@app.route('/html_format', methods=['POST'])
@login_required  # 需要登录才能访问
def clean_html():
    try:
        content = request.json.get('content')
        
        return jsonify({'result': clean(content)})
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

import autopep8
import black
# 选择题格式化
def to_choose(content):
    # 解析HTML内容
    content = clean(content)
    soup = BeautifulSoup(content, 'lxml')
    plain_text = soup.get_text(separator='\n')  # 获取纯文本
    # 使用正则表达式分割正文和ABCD选项
    pattern = r"^(.*?)\s*(A[ .]?.*?)\s*(B[ .]?.*?)\s*(C[ .]?.*?)\s*(D[ .]?.*?)$"
    match = re.match(pattern, plain_text, re.DOTALL)

    if not match:
        return False

    # 提取正文和选项
    body_text, option_a, option_b, option_c, option_d = match.groups()
    if "（▲）" not in body_text:
        body_text = re.sub(r'[\(（]\s*[\)）]$', '（▲）', body_text)
        if "（▲）" not in body_text:
            body_text += "（▲）"
    # 构建新的HTML结构
    formatted_soup = BeautifulSoup("", 'lxml')
    body_div = formatted_soup.new_tag('div')
    # 创建正文的div
    for i in body_text.split("\n"):
        line_div  = formatted_soup.new_tag('div')
        line_div .string = i
        body_div.append(line_div)
    # 将 body_div 添加到 formatted_soup
    formatted_soup.append(body_div)
    # 查找并插入图片
    img_tag = soup.find('img')
    if img_tag:
        formatted_soup.append(img_tag)
        # 创建选项的div
    for option in [option_a, option_b, option_c, option_d]:
        option_div = formatted_soup.new_tag('div')
        option_div.string = option.strip()
        formatted_soup.append(option_div)
    # 返回仅包含div和图片的结构
    return str(formatted_soup)

@app.route('/choose_format', methods=['POST'])
@login_required  # 需要登录才能访问
def clean_choose():
    try:
        content = request.json.get('content')
        content = clean(content)

        # 返回仅包含div和图片的结构
        return jsonify({'result': to_choose(content)})

    except Exception as e:
        return jsonify({'error': str(e)}), 500

# GPT模块----------------------------------------------------------------------------------------
from sparkai.llm.llm import ChatSparkLLM, ChunkPrintHandler
from sparkai.core.messages import ChatMessage
from openai import OpenAI
# 配置讯飞API用户信息
#星火认知大模型Spark Max的URL值，其他版本大模型URL值请前往文档（https://www.xfyun.cn/doc/spark/Web.html）查看
SPARKAI_URL = app.config['SPARKAI_URL']
#星火认知大模型调用秘钥信息，请前往讯飞开放平台控制台（https://console.xfyun.cn/services/bm35）查看
SPARKAI_APP_ID = app.config['SPARKAI_APP_ID']
SPARKAI_API_SECRET = app.config['SPARKAI_API_SECRET']
SPARKAI_API_KEY = app.config['SPARKAI_API_KEY']
#星火认知大模型Spark Max的domain值，其他版本大模型domain值请前往文档（https://www.xfyun.cn/doc/spark/Web.html）查看
SPARKAI_DOMAIN = 'generalv3.5'

@app.route('/send', methods=['POST'])
@login_required  # 确保用户已登录
def send_message():
    client = OpenAI(api_key="sk-75e3bd294770479cbf83bd3b3aa2d2ae", base_url="https://api.deepseek.com")
    user_input = request.form.get('user_input')
    remaining = 5 - current_user.request_count
    if not user_input:
        return jsonify({'response': '未找到用户输入！', 'remaining': remaining}), 400

    user = current_user  # 获取当前用户

    today = date.today()
    if user.last_request_date != today:
        # 如果上次请求日期不是今天，重置请求次数
        user.request_count = 0
        user.last_request_date = today

    if user.request_count >= 5 and user.role != 'admin':
        return jsonify({
            'response': '您今天的请求次数已达到上限，请明天再试。',
            'remaining': 0
        }), 403  # 403 Forbidden

    try:
        # 构建消息列表，包含用户的提问
        promote1 = "你是一位python编程小助手。1. 你只回答与编程相关的问题。如果问题不涉及编程，你需回复：'××'。2.代码务必加上注释并且给出算法思路和编写解释。3. 输出代码越初级越好，除非用户要求，否则禁止用自定义函数以及格式化输出。4. 避免用try、模块使用、内置函数。"
        promote2 = "你是一位小助手，可以回答任何问题。"
        messages=[
                {"role": "system", "content": promote1},
                {"role": "user", "content": user_input}
          ]
        # 调用 Deepseek 的聊天完成接口（使用 Chat API）
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=messages,
            max_tokens=1024,
            temperature=0.7,
            stream=False
        )

        # 提取生成的文本
        assistant_message = response.choices[0].message.content

        if '××' in assistant_message.lower():
            assistant_message = '抱歉，请勿提出无关的问题，本次回答不计次数（如果造成了误判，请更换提问方式或者更详细地提问）。'
        else:
            # 增加请求次数
            user.request_count += 1
            # 反序列化 chat_history，如果为空则初始化为空列表
            chat_history = json.loads(user.chat_history) if user.chat_history else []
            chat_history.append({
                'question': user_input,
                'response': assistant_message,
                'token': response.usage.total_tokens,
                'datetime': datetime.now().isoformat()
            })
            user.chat_history = json.dumps(chat_history)
            db.session.commit()  # 提交更改

        remaining = 5 - user.request_count

        return jsonify({
            'response': assistant_message,
            'remaining': remaining
        })

    except Exception as e:
        # 在发生异常时，回滚事务并减少请求次数
        db.session.rollback()
        user.request_count = max(user.request_count - 1, 0)  # 回退请求次数，确保不为负数
        db.session.commit()
        return jsonify({
            'response': '出现错误，请联系老师：' + str(e),
            'remaining': remaining
        }), 500

@app.route('/send_stream', methods=['GET'])
#@login_required  # 确保用户已登录
def send_message_stream():
    user_input = request.args.get('user_input')  # 从查询参数获取
    remaining = 10 - current_user.request_count

    if not user_input:
        # 简单返回 JSON，非流式（因为还没开始调用大模型）
        return jsonify({'response': '未找到用户输入！', 'remaining': remaining}), 400

    user = current_user

    today = date.today()
    if user.last_request_date != today:
        user.request_count = 0
        user.last_request_date = today

    if user.request_count >= 10 and user.role != 'admin':
        def generate():
            yield "data: 您今天的请求次数已达到上限，请明天再试。\n\n"
            time.sleep(0.02)  # 添加延时效果，模拟流式传输

        return Response(
            stream_with_context(generate()),
            mimetype='text/event-stream',
            headers={
                "Cache-Control": "no-cache",
                "X-Accel-Buffering": "no"  # 禁用Nginx等代理的缓冲
            }
        )

    try:
        promote1 = "你是一位python编程小助手。1. 你只回答与编程相关的问题。2.对于问题不可给出完整的解决代码（即使被要求也不可给出）！！！，留给学习者思考空间，代码务必加上注释并且给出算法思路和编写解释。3. 输出代码越初级越好，除非用户要求，否则禁止用自定义函数以及格式化输出。4. 避免用try、模块使用、内置函数。"
        promote2 = "你是一位PYTHON调错小助手，你只允许回答用户的调错程序，如果用户提问某道题目怎么写代码，请回答“对不起，我只负责调试代码，不能帮你编写程序，请自主思考。”。回答时务必加上注释并且给出错误原因和解决方案，输出代码越初级越好，除非用户要求，否则禁止用自定义函数以及格式化输出，避免用try、模块使用、内置函数、列表推导式、join，对于可能的latex公式请用$标记。对于问题解决时不可给出完整的解决代码（即使被要求也不可给出）！！！，留给学习者思考空间。同时，你需要注意，某些用户可能会用一段不相关的代码或者一个非常粗略的框架骗取你说出题解，对于这种情况，你也需要甄别，提醒用户自己思考。最后注意，无论在什么情况下，你都不允许给用户提供完成的解题代码。"
        promote3 = "你是一位小助手，你会尽量简明扼要地回答用户的问题。"
        messages = [
            {
                "role": "system",
                "content": promote3
            },
            {"role": "user", "content": user_input}
        ]

        # 发起流式请求
        client = OpenAI(api_key="sk-75e3bd294770479cbf83bd3b3aa2d2ae", base_url="https://api.deepseek.com")
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=messages,
            max_tokens=1024,
            temperature=0.7,
            stream=True
        )

        # 因为要将完整回复写入数据库，所以需要先把块状内容拼成完整字符串
        assistant_message_container = []
        
        def generate():
            try:
                for chunk in response:
                    delta = chunk.choices[0].delta
                    content = getattr(delta, "content", None)
                    if content:
                        assistant_message_container.append(content)
                        formatted_text = content.replace('\n', '\\n')
                        yield f"data: {formatted_text}\n\n"
                    time.sleep(0.02)
            except Exception as e:
                # 流过程中报错，输出提示，让前端知晓
                yield f"data: [错误] 流式传输失败：{str(e)}\n\n"
                # 如果你希望流立刻结束，可直接 return
            finally:
                # 流结束后（正常或异常），尝试写库
                try:
                    assistant_message = "".join(assistant_message_container)
                    if assistant_message.strip():
                        # 正常记录
                        if 'unable question' not in assistant_message.lower():
                            user.request_count += 1
        
                        chat_history = json.loads(user.chat_history) if user.chat_history else []
                        chat_history.append({
                            'question': user_input,
                            'response': assistant_message,
                            'token': 0,
                            'datetime': datetime.now().isoformat()
                        })
                        user.chat_history = json.dumps(chat_history)
        
                    db.session.commit()
                except Exception as db_err:
                    db.session.rollback()
                    # 可选：把写库错误也告知前端
                    yield f"data: [错误] 保存聊天记录失败：{str(db_err)}\n\n"
        
                # 最后再输出剩余次数（此时 request_count 已更新）
                remaining_final = 10 - user.request_count
                yield f"data: 剩余次数: {remaining_final}\n\n"
        
        resp = Response(
            stream_with_context(generate()),
            mimetype='text/event-stream',
            headers={
                "Cache-Control": "no-cache",
                "X-Accel-Buffering": "no"
            }
        )
        return resp

    except Exception as e:
        db.session.rollback()
        user.request_count = max(user.request_count - 1, 0)
        db.session.commit()
        return jsonify({
            'response': '出现错误，请联系老师：' + str(e),
            'remaining': remaining
        }), 500



@app.route('/send2', methods=['POST'])
def send_message2():
    user_input = request.form.get('user_input')
    remaining = 10 - current_user.request_count
    if not user_input:
        return jsonify({'response': '未找到用户输入！', 'remaining': remaining}), 400

    user = current_user  # 获取当前用户

    today = date.today()
    if user.last_request_date != today:
        # 如果上次请求日期不是今天，重置请求次数
        user.request_count = 0
        user.last_request_date = today

    if user.request_count >= 10 and user.role != 'admin':
        return jsonify({
            'response': '您今天的请求次数已达到上限，请明天再试。',
            'remaining': 0
        }), 403  # 403 Forbidden

    try:
        # 处理用户输入
        question = "1.注意！你只回答与python编程相关的问题。如果问题不涉及编程，你需回复：'unable question'。2.代码务必加上注释并且给出算法思路和编写解释。3.输出代码越初级越好，除非用户要求，否则禁止用自定义函数以及格式化输出。4.避免用try、模块使用、内置函数。5.以下是问题：" + user_input

        spark = ChatSparkLLM(
            spark_api_url=SPARKAI_URL,
            spark_app_id=SPARKAI_APP_ID,
            spark_api_key=SPARKAI_API_KEY,
            spark_api_secret=SPARKAI_API_SECRET,
            spark_llm_domain=SPARKAI_DOMAIN,
            streaming=False
        )
        #以下为星火GPT的调用
        
        handler = ChunkPrintHandler()
        messages = [ChatMessage(
            role="user",
            content=question
        )]
        
        result = spark.generate([messages], callbacks=[handler])
        response = result.generations[0][0].text
        
        #ChatGPT
        #response = chat(question)
        
        if 'unable question' in response or 'Unable question' in response:
            response = '抱歉，我只回答编程相关的问题，本次回答不计次数（如果造成了误判，请更换提问方式或者更详细地提问）。'
        else:
            # 增加请求次数
            user.request_count += 1
            # 反序列化 chat_history，如果为空则初始化为空列表
            chat_history = json.loads(user.chat_history) if user.chat_history else []
            chat_history.append({'response': response, 'token': result.llm_output['token_usage']['total_tokens'], 'datetime': datetime.now().isoformat()})
            user.chat_history = json.dumps(chat_history)
            db.session.commit()  # 提交更改
            
        remaining = 10 - user.request_count
        
        return jsonify({
            'response': response,
            'remaining': remaining
        })
    except Exception as e:
        # 在发生异常时，回滚事务并减少请求次数
        db.session.rollback()
        user.request_count -= 1  # 回退请求次数
        db.session.commit()
        return jsonify({
            'response': '出现错误，请联系老师：' + str(e) + str(response),
            'remaining': remaining
        }), 500

# DeepSeek生成解析
import requests

@app.route('/generate_analysis/<gpttype>', methods=['POST'])
@login_required  # 需要登录才能访问
def generate_analysis(gpttype):
    try:
        content = request.json.get('content')
        answer = request.json.get('answer')
        prompt = request.json.get('prompt')

        if gpttype == 'baidu':
            result = call_baidu_gpt_api(prompt+"题目："+content+"。答案：" + answer)
            return jsonify({'analysis': result})
        elif gpttype == "deepseekv3":
            client = OpenAI(api_key="sk-75e3bd294770479cbf83bd3b3aa2d2ae", base_url="https://api.deepseek.com")
            # 构建消息列表，包含用户的提问
            messages=[
                    {"role": "system", "content": prompt},
                    {"role": "user", "content": "题目：" + content + "。答案：" + answer}
              ]
            # 调用 Deepseek 的聊天完成接口（使用 Chat API）
            response = client.chat.completions.create(
                model="deepseek-chat",
                messages=messages,
                max_tokens=1024,
                temperature=0.7,
                stream=False
            )
            # 提取生成的文本
            generated_analysis = response.choices[0].message.content
            return jsonify({'analysis': generated_analysis})
        elif gpttype == "chatgpt3.5":
            return jsonify({'analysis': chatgpt(prompt + "题目：" + content + "。答案：" + answer, "gpt-3.5-turbo")})
        elif gpttype == "chatgpt4omini":
            return jsonify({'analysis': chatgpt(prompt + "题目：" + content + "。答案：" + answer, "gpt-4o-mini")})
        elif gpttype == "chatgpt4.5":
            #return jsonify({'analysis': '暂未开通'})
            return jsonify({'analysis': chatgpt(prompt + "题目：" + content + "。答案：" + answer, "gpt-4.5-preview-2025-02-27")})
        elif gpttype == "chatgpt4o":
            #return jsonify({'analysis': '暂未开通'})
            return jsonify({'analysis': chatgpt(prompt + "题目：" + content + "。答案：" + answer, "gpt-4o")})
        elif gpttype == "chatgpto1mini":
            #return jsonify({'analysis': '暂未开通'})
            return jsonify({'analysis': chatgpt(prompt + "题目：" + content + "。答案：" + answer, "o1-mini-2024-09-12")})
        elif gpttype == "chatgpto1":
            #return jsonify({'analysis': '暂未开通'})
            return jsonify({'analysis': chatgpt(prompt + "题目：" + content + "。答案：" + answer, "o1")})
        elif gpttype == "chatgpto3mini":
            #return jsonify({'analysis': '暂未开通'})
            return jsonify({'analysis': chatgpt(prompt + "题目：" + content + "。答案：" + answer, "o3-mini")})
        elif gpttype == "chatgpt4":
            #return jsonify({'analysis': '暂未开通'})
            return jsonify({'analysis': chatgpt(prompt + "题目：" + content + "。答案：" + answer, "gpt-4")})
        elif gpttype == "deepseekr1":
            client = OpenAI(api_key="sk-75e3bd294770479cbf83bd3b3aa2d2ae", base_url="https://api.deepseek.com")
            # 构建消息列表，包含用户的提问
            messages=[
                    {"role": "system", "content": prompt},
                    {"role": "user", "content": "题目：" + content + "。答案：" + answer}
              ]
    
            # 调用 Deepseek 的聊天完成接口（使用 Chat API）
            response = client.chat.completions.create(
                model="deepseek-reasoner",
                messages=messages,
                max_tokens=1024,
                temperature=0.7,
                stream=False
            )
    
            # 提取生成的文本
            generated_analysis = response.choices[0].message.content
            return jsonify({'analysis': generated_analysis})
    except Exception as e:
        return jsonify({'error': str(e)}), 500



# 图片上传API（控件用）----------------------------------------------------------------------------------------
# 生成随机文件名
def generate_random_string(length):
    characters = 'ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789'
    return ''.join(random.choice(characters) for i in range(length))

# 图片上传
@app.route('/upload_image', methods=['POST'])
@login_required  # 需要登录才能访问
def upload_image():
    if 'file' not in request.files:
        return jsonify({'error': '无文件'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '未选择文件'}), 400
    if file:
        _, file_extension = os.path.splitext(file.filename)
        random_filename = generate_random_string(16) + file_extension
        filepath = os.path.join('./uploads', random_filename)
        file.save(filepath)
        return jsonify({'url': '_uploads/photos/' + random_filename}), 200
    return jsonify({'error': '上传失败'}), 500

# 文件上传API（云盘）----------------------------------------------------------------------------------------
# 文件上传
@app.route('/upload', methods=['POST'])
@login_required  # 需要登录才能访问
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': '无文件'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '未选择文件'}), 400
    if file:
        filename = secure_filename_with_chinese(file.filename)
        filepath = os.path.join('./uploads/files', filename)
        file.save(filepath)
        return jsonify({'message': '上传成功'}), 200
    return jsonify({'error': '上传失败'}), 500

# 获取文件列表
@app.route('/api/files')
@login_required  # 需要登录才能访问
def list_files():
    directory = './uploads/files'
    files = []
    for filename in os.listdir(directory):
        if os.path.isfile(os.path.join(directory, filename)):
            stat = os.stat(os.path.join(directory, filename))
            size = stat.st_size
            mtime = datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M:%S')
            files.append({
                'filename': filename,
                'size': size,
                'mtime': mtime
            })
    files.sort(key=lambda x: x['mtime'], reverse=True)
    return jsonify(files)

# 选中题目打包下载API----------------------------------------------------------------------------------------
@app.route('/generate_html', methods=['POST'])
@login_required
def generate_html_endpoint():
    try:
        question_ids = request.json.get('question_ids', [])
        list_name = request.json.get('list_name', '')
    
        if not question_ids:
            return jsonify({'error': '未提供题目ID'}), 400
    
        order = case(
            {id: index for index, id in enumerate(question_ids)}, 
            value=Question.id
        )
        questions = Question.query.filter(Question.id.in_(question_ids)).order_by(order).all()
        
        if list_name:
            current_user_id = current_user.get_id()
            download_record = DownloadRecord(user_id=current_user_id, list_value=','.join(map(str, question_ids)), list_name=list_name)
            db.session.add(download_record)
            db.session.commit()
    
        html_content = generate_html(questions)
        return html_content
    except Exception as e:
        return str(e)

# 生成HTML内容
def generate_html(questions):
    html_content = """
<html>
<head>
<style>
.q-line { display: flex; align-items: flex-start; }
.q-num { margin-right: 4px; white-space: nowrap; }
</style>
</head>
<body>
"""
    num = 1
    for question in questions:
        html_content += (
            f'<div class="q-line">'
            f'<span class="q-num">{num}.</span>'
            f'{question.content}'
            f'</div>'
        )
        num += 1

    num = 1
    html_content += "<br>【答案与解析】<br>"
    for question in questions:
        html_content += f'{num}. {question.answer}<br>'
        if question.analysis:
            html_content += question.analysis + '<br>'
        num += 1

    html_content += "</body></html>"
    return html_content

# 保持前端返回的题目ID的顺序
from sqlalchemy.sql.expression import case
from sqlalchemy.orm.exc import NoResultFound
@app.route('/download_questions', methods=['POST'])
@login_required  # 需要登录才能访问
def download_questions():
    question_ids = request.json.get('question_ids', [])
    list_name = request.json.get('list_name', '')
    
    if not question_ids:
        return jsonify({'error': '未提供题目ID'}), 400
        
    log_request_info('download_questions', current_user, '下载题目'+','.join([str(i) for i in question_ids]))
    order = case(
        {id: index for index, id in enumerate(question_ids)}, 
        value=Question.id
    )
    
    questions = Question.query.filter(Question.id.in_(question_ids)).order_by(order).all()
    
    if list_name != "":
        current_user_id = current_user.get_id()
        download_record = DownloadRecord(user_id=current_user_id, list_value=','.join([str(i) for i in question_ids]), list_name=list_name)
        db.session.add(download_record)
        db.session.commit()
    
    html_content = generate_html(questions)

    try:
        output_docx = pypandoc.convert_text(
            html_content, 
            to='docx', 
            format='html', 
            extra_args=[
                '--reference-doc', './tmp_doc/Doc1.docx'
            ],
            outputfile="./tmp_doc/output.docx"
        )
    except Exception as e:
        print("Error during conversion:"+str(e))

    with open("./tmp_doc/output.docx", "rb") as file:
        docx_data = file.read()

    response = Response(docx_data, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response.headers["Content-Disposition"] = "attachment; filename=questions.docx"
    return response

# 保存配置文件----------------------------------------------------------------------------------------
@app.route('/save_config', methods=['POST'])
def save_config():
    data = request.json
    

    # 更新全局变量或配置文件
    config_data = {
        "ALLOWED_IPS": data.get('ALLOWED_IPS'),
        "BAIDU_GPT_ID": data.get('BAIDU_GPT_ID'),
        "BAIDU_GPT_SECRET_KEY": data.get('BAIDU_GPT_SECRET_KEY'),
        "CHATGPT_KEY": data.get('CHATGPT_KEY'),
        "CHATGPT_BASE_URL": data.get('CHATGPT_BASE_URL'),
        "SPARKAI_URL": data.get('SPARKAI_URL'),
        "SPARKAI_APP_ID": data.get('SPARKAI_APP_ID'),
        "SPARKAI_API_SECRET": data.get('SPARKAI_API_SECRET'),
        "SPARKAI_API_KEY": data.get('SPARKAI_API_KEY'),
    }
    with open('config.json', 'w') as config_file:
        json.dump(config_data, config_file, indent=4, ensure_ascii=False)
    # 重新加载配置并更新到 app.config
    app.config.update(config_data)
    load_config()
    return jsonify({"message": "配置已成功保存"})

# word处理功能，用于tamcat网站----------------------------------------------------------------------------------------
DOWNLOAD_DOC_FOLDER = './tomcat_doc/download_docx'
UPLOAD_DOC_FOLDER = './tomcat_doc/upload_doc'
app.config['DOWNLOAD_DOC_FOLDER'] = DOWNLOAD_DOC_FOLDER
app.config['UPLOAD_DOC_FOLDER'] = UPLOAD_DOC_FOLDER
if not os.path.exists(UPLOAD_DOC_FOLDER):
    os.makedirs(UPLOAD_DOC_FOLDER)
if not os.path.exists(DOWNLOAD_DOC_FOLDER):
    os.makedirs(DOWNLOAD_DOC_FOLDER)

# doc上传页面
@app.route('/upload_docx')
@login_required  # 需要登录才能访问
def upload_docx_page():
    return render_template('upload_docx.html')

# 上传docx文件并处理
@app.route('/upload_docx', methods=['POST'])
@login_required  # 需要登录才能访问
def upload_docx():
    if request.method == 'POST':
        file = request.files['file']
        if file:
            filepath = os.path.join(app.config['UPLOAD_DOC_FOLDER'], file.filename)
            file.save(filepath)
            process_file(filepath, file.filename)
            return send_from_directory(app.config['DOWNLOAD_DOC_FOLDER'], file.filename, as_attachment=True)
    return "Failed"

# 处理上传的docx文件
def process_file(input_path, filename):
    file_name, file_extension = os.path.splitext(input_path)
    if file_extension.lower() == '.doc':
        os.system(f"unoconv -f docx '{input_path}'")
        input_path = f"{file_name}.docx"
    docx = Document(input_path)
    for paragraph in docx.paragraphs:
        for run in paragraph.runs:
            run.text = run.text.replace('??', '\n')
            run.text = run.text.replace('&nbsp;', ' ')
            run.text = run.text.replace('&gt;', '>')
            run.text = run.text.replace('&lt;', '<')
            run.text = run.text.replace('&lsquo;', "'")
            run.text = run.text.replace('&rsquo;', "'")
            run.text = run.text.replace('&amp;', "&")
            run.text = run.text.replace('&quot;', '"')
            run.text = run.text.replace('?', '\n')
            run.text = run.text.replace('&ldquo;', '"')
            run.text = run.text.replace('&rdquo;', '"')
            
    output_path = os.path.join(app.config['DOWNLOAD_DOC_FOLDER'], filename)
    docx.save(output_path)

def check_ip(ip1, ip2):
    ip1 = ip1.split('.')
    x = 0
    for i in ip2.split('.'):
        if ip1[x]!=i and i!='*':
            return False
        x += 1
    return True

def check_allowed_ip(ip, ALLOWED_IPS):
    for ips in ALLOWED_IPS:
        if check_ip(ip, ips):
            return True
    return False
def get_ip_location(ip):
    try:
        response = requests.get(f"http://ip-api.com/json/{ip}")
        ip_data = response.json()
        # 检查请求是否成功
        if ip_data.get("status") == "success":
            return f"{ip_data['country']}, {ip_data['org']}"
        else:
            return "Location not found"
    except Exception as e:
        return "Location not found：" + str(e)
#操作记录函数
def log_request_info(page, user, extra_info=""):
    # 获取 IP 地址
    ip = request.headers.get('X-Real-IP', request.remote_addr).split(',')[0].strip()
    if check_allowed_ip(ip, app.config['ALLOWED_IPS']):
        return

    # 获取地理位置
    location = ""  # get_ip_location(ip)
    access_time = datetime.now()
    
    # 根据用户认证状态获取用户信息
    if user.is_authenticated:
        user_id = user.id
        username = user.username  # 假设用户模型有username字段
    else:
        user_id = None
        username = "Anonymous"  # 匿名用户使用默认名称

    # 检查是否在1秒内有相同记录
    time_threshold = access_time - timedelta(seconds=1)
    existing_log = None

    if user.is_authenticated:
        # 已登录用户：基于user_id和页面检查
        existing_log = RouteLog.query.filter_by(
            user_id=user_id, 
            page=page
        ).filter(
            RouteLog.access_time >= time_threshold
        ).first()
    else:
        # 匿名用户：基于IP和页面检查
        existing_log = RouteLog.query.filter_by(
            ip_address=ip, 
            page=page
        ).filter(
            RouteLog.access_time >= time_threshold
        ).first()

    if existing_log:
        return  # 存在重复记录，不保存

    # 创建并保存新记录
    route_log = RouteLog(
        user_id=user_id,
        user_name=username,
        ip_address=ip,
        access_time=access_time,
        page=page,
        location=location,
        extra_info=extra_info
    )
    db.session.add(route_log)
    db.session.commit()